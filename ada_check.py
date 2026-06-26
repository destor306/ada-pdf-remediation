#!/usr/bin/env python3
"""
ADA Quality Checker — free, local, no API required.

Runs a layered check stack against a remediated .docx and/or exported .pdf:

  Layer 1 — Docx structure audit (always runs, pure Python)
    • Heading hierarchy (no skipped levels)
    • Table header rows present
    • Alt text coverage for images
    • Text coverage vs. source PDF (detect content loss)
    • Language declaration

  Layer 2 — PDF tag-tree audit via pikepdf (always runs if a PDF is given)
    • Reads the same StructTreeRoot data Acrobat's Full Check reads —
      no Java, no install, no internet required
    • Tagged-PDF flag, document language, title, DisplayDocTitle
    • Heading tag hierarchy in the actual PDF (not just the docx)
    • Table TH cells and /Scope attribute
    • Figure tags and their /Alt text
    • Form field tagging and tooltip (/TU) presence

  Layer 3 — PDF/UA validation via VeraPDF (runs if Java + VeraPDF available)
    • Full PDF/UA-1 (ISO 14289-1) conformance check
    • Reports clause-level failures

Usage:
  python3 ada_check.py source.pdf remediated.docx [remediated.pdf]

Install VeraPDF (optional but recommended for production):
  sudo apt install default-jre
  wget https://github.com/veraPDF/veraPDF-apps/releases/download/v1.26.2/verapdf-greenfield-1.26.2-installer.zip
  unzip verapdf-greenfield-*.zip && sudo bash verapdf-greenfield*/verapdf_install.sh
"""

import sys
import os
import re
import subprocess
import tempfile
import json
from pathlib import Path
from dataclasses import dataclass, field

sys.stdout.reconfigure(encoding='utf-8')

import pdfplumber
import pikepdf
from pdf2image import convert_from_path
from docx import Document
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Result types
# ---------------------------------------------------------------------------

@dataclass
class Issue:
    severity: str   # "error" | "warning" | "info"
    category: str
    message: str

@dataclass
class CheckReport:
    source_pdf: str
    docx_path: str
    pdf_path: str | None = None
    issues: list[Issue] = field(default_factory=list)
    passed: int = 0
    failed: int = 0

    def add(self, severity: str, category: str, message: str):
        self.issues.append(Issue(severity, category, message))
        if severity == "error":
            self.failed += 1
        else:
            self.passed += 1

    def ok(self, category: str, message: str):
        self.add("info", category, f"✓ {message}")
        self.passed += 1

    def warn(self, category: str, message: str):
        self.add("warning", category, f"⚠ {message}")

    def error(self, category: str, message: str):
        self.add("error", category, f"✗ {message}")
        self.failed += 1

    def print_report(self):
        print(f"\n{'='*60}")
        print(f"ADA Quality Report")
        print(f"  Source:  {self.source_pdf}")
        print(f"  Output:  {self.docx_path}")
        if self.pdf_path:
            print(f"  PDF:     {self.pdf_path}")
        print(f"{'='*60}")

        # Group by category
        categories = {}
        for issue in self.issues:
            categories.setdefault(issue.category, []).append(issue)

        for cat, items in categories.items():
            print(f"\n[{cat}]")
            for item in items:
                print(f"  {item.message}")

        print(f"\n{'='*60}")
        errors = [i for i in self.issues if i.severity == "error"]
        warnings = [i for i in self.issues if i.severity == "warning"]
        print(f"Result: {len(errors)} errors, {len(warnings)} warnings")
        if not errors:
            print("Status: PASS (no blocking errors)")
        else:
            print("Status: FAIL — fix errors before publishing")
        print(f"{'='*60}\n")
        return len(errors) == 0


# ---------------------------------------------------------------------------
# Layer 1: Docx structure audit
# ---------------------------------------------------------------------------

def _iter_block_items(container):
    """Yield paragraphs and tables in true document order (as they appear in
    the XML body), instead of python-docx's separate .paragraphs/.tables
    lists which lose interleaving between the two."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    parent_elm = container.element.body if hasattr(container, "element") else container._tc
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, container)
        elif child.tag == qn("w:tbl"):
            yield Table(child, container)


def _iter_all_paragraphs(container):
    """Yield every paragraph in a Document (or table cell) in true reading
    order, including ones nested inside table cells. pdf2docx renders most
    form layouts as tables, so headings routinely end up inside cells —
    doc.paragraphs alone misses them, and a naive paragraphs-then-tables
    traversal would scramble the order needed for level-skip checks."""
    from docx.table import Table

    for block in _iter_block_items(container):
        if isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    yield from _iter_all_paragraphs(cell)
        else:
            yield block


def check_heading_hierarchy(doc: Document, report: CheckReport):
    """Headings must not skip levels (e.g., H1 → H3 without H2)."""
    heading_levels = []
    for para in _iter_all_paragraphs(doc):
        style = para.style.name
        if style.startswith("Heading "):
            try:
                level = int(style.split()[-1])
                heading_levels.append(level)
            except ValueError:
                pass

    if not heading_levels:
        report.error("Headings", "No heading styles found — document has no structure")
        return

    report.ok("Headings", f"Found {len(heading_levels)} headings")

    if heading_levels[0] != 1:
        report.warn("Headings", f"Document does not start with Heading 1 (starts at H{heading_levels[0]})")

    for i in range(1, len(heading_levels)):
        prev, curr = heading_levels[i - 1], heading_levels[i]
        if curr > prev + 1:
            report.error("Headings", f"Skipped heading level: H{prev} → H{curr} (missing H{prev+1})")


def check_table_headers(doc: Document, report: CheckReport):
    """Every table must have at least one header row (w:tblHeader)."""
    tables = doc.tables
    if not tables:
        report.ok("Tables", "No tables in document")
        return

    report.ok("Tables", f"Found {len(tables)} table(s)")
    missing_headers = 0

    for t_idx, table in enumerate(tables, 1):
        has_header = False
        for row in table.rows:
            trPr = row._tr.find(qn("w:trPr"))
            if trPr is not None and trPr.find(qn("w:tblHeader")) is not None:
                has_header = True
                break
        if not has_header:
            report.error("Tables", f"Table {t_idx}: no header row (w:tblHeader) — screen readers can't identify column headers")
            missing_headers += 1

    if missing_headers == 0:
        report.ok("Tables", "All tables have header rows marked")


def check_alt_text(doc: Document, report: CheckReport):
    """Inline images should have alt text (wp:docPr descr attribute)."""
    ns = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }
    inline_images = doc.element.body.findall(".//wp:inline", ns)
    if not inline_images:
        report.ok("Alt Text", "No inline images found")
        return

    missing = 0
    for img in inline_images:
        docPr = img.find("wp:docPr", ns)
        descr = docPr.get("descr", "").strip() if docPr is not None else ""
        if not descr:
            missing += 1

    if missing:
        report.error("Alt Text", f"{missing}/{len(inline_images)} image(s) missing alt text (descr attribute)")
    else:
        report.ok("Alt Text", f"All {len(inline_images)} image(s) have alt text")


def check_language(doc: Document, report: CheckReport):
    """Document language should be declared."""
    lang = doc.core_properties.language
    if lang:
        report.ok("Language", f"Document language set: {lang}")
    else:
        report.error("Language", "Document language not declared (required for screen readers)")


def check_text_coverage(source_pdf: str, doc: Document, report: CheckReport):
    """
    Compare word coverage between source PDF and output docx.
    Flags if more than 10% of source words are missing from the output.
    """
    # Extract source words
    source_words = set()
    with pdfplumber.open(source_pdf) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            for w in text.split():
                source_words.add(w.lower().strip(".,;:()[]\"'"))

    # Extract output words
    output_words = set()
    for para in doc.paragraphs:
        for w in para.text.split():
            output_words.add(w.lower().strip(".,;:()[]\"'"))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for w in cell.text.split():
                    output_words.add(w.lower().strip(".,;:()[]\"'"))

    # Filter out very short words (noise)
    source_words = {w for w in source_words if len(w) > 2}
    output_words = {w for w in output_words if len(w) > 2}

    if not source_words:
        report.warn("Coverage", "Could not extract text from source PDF for comparison")
        return

    missing = source_words - output_words
    coverage = 1.0 - len(missing) / len(source_words)
    pct = coverage * 100

    if pct >= 90:
        report.ok("Coverage", f"Text coverage: {pct:.1f}% ({len(source_words)} source words)")
    elif pct >= 75:
        report.warn("Coverage", f"Text coverage: {pct:.1f}% — some content may be missing")
    else:
        report.error("Coverage", f"Text coverage: {pct:.1f}% — significant content loss detected")

    if missing and len(missing) <= 20:
        report.warn("Coverage", f"Missing words sample: {', '.join(sorted(missing)[:20])}")


def check_empty_headings(doc: Document, report: CheckReport):
    """Headings must not be empty."""
    empty = 0
    for para in _iter_all_paragraphs(doc):
        if para.style.name.startswith("Heading ") and not para.text.strip():
            empty += 1
    if empty:
        report.error("Headings", f"{empty} empty heading paragraph(s) found")
    else:
        report.ok("Headings", "No empty headings")


def run_docx_checks(source_pdf: str, docx_path: str, report: CheckReport):
    print("Running docx structure audit...")
    doc = Document(docx_path)
    check_language(doc, report)
    check_heading_hierarchy(doc, report)
    check_empty_headings(doc, report)
    check_table_headers(doc, report)
    check_alt_text(doc, report)
    check_text_coverage(source_pdf, doc, report)


# ---------------------------------------------------------------------------
# Layer 2: PDF tag-tree audit (pikepdf) — mirrors what Acrobat's Full Check
# reads directly off the PDF's StructTreeRoot. No Java, no install needed.
# ---------------------------------------------------------------------------

def _walk_struct(node, callback):
    """Recursively visit structure elements under `node`, calling
    callback(tag_name, elem) for every real StructElem (skips MCID ints
    and /MCR /OBJR marked-content references, which are leaves)."""
    if isinstance(node, pikepdf.Array):
        for child in node:
            _walk_struct(child, callback)
        return
    if not isinstance(node, pikepdf.Dictionary):
        return
    if str(node.get("/Type", "")) in ("/MCR", "/OBJR"):
        return
    tag = node.get("/S")
    tag_name = str(tag)[1:] if tag is not None else None
    if tag_name:
        callback(tag_name, node)
    kids = node.get("/K")
    if kids is not None:
        _walk_struct(kids, callback)


def _struct_attr(elem, name):
    """Look up an attribute (e.g. /Scope) across an element's /A attribute
    object(s), which may be a single dict or an array of dicts."""
    attrs = elem.get("/A")
    if attrs is None:
        return None
    attr_list = attrs if isinstance(attrs, pikepdf.Array) else [attrs]
    for a in attr_list:
        if isinstance(a, pikepdf.Dictionary) and a.get(name) is not None:
            return a.get(name)
    return None


def check_pdf_tags(pdf_path: str, report: CheckReport):
    """Inspect the actual exported PDF's structure tree — the same data
    Acrobat's accessibility checker reads — rather than the pre-export docx."""
    print("Running PDF tag-tree audit (pikepdf)...")
    try:
        pdf = pikepdf.open(pdf_path)
    except Exception as e:
        report.warn("PDF Tags", f"Could not open PDF for tag inspection: {e}")
        return

    root = pdf.Root

    mark_info = root.get("/MarkInfo")
    tagged = bool(mark_info and mark_info.get("/Marked"))
    if tagged:
        report.ok("PDF Tags", "Document is marked as Tagged PDF (/MarkInfo /Marked)")
    else:
        report.error("PDF Tags", "Document is NOT marked as Tagged PDF — Acrobat will flag this immediately")

    lang = root.get("/Lang")
    if lang:
        report.ok("PDF Tags", f"Document language set in PDF catalog: {lang}")
    else:
        report.error("PDF Tags", "No /Lang set in PDF catalog (required for screen readers)")

    title = None
    try:
        title = pdf.docinfo.get("/Title")
    except Exception:
        pass
    if title:
        report.ok("PDF Tags", f"Document title set: {title}")
    else:
        report.error("PDF Tags", "No document title set (Document Properties > Title)")

    viewer_prefs = root.get("/ViewerPreferences")
    if viewer_prefs and viewer_prefs.get("/DisplayDocTitle"):
        report.ok("PDF Tags", "DisplayDocTitle is set — title bar will show document title")
    else:
        report.warn("PDF Tags", "DisplayDocTitle not set true — title bar may show filename instead of title")

    struct_root = root.get("/StructTreeRoot")
    if not struct_root:
        report.error("PDF Tags", "No /StructTreeRoot found — document has no logical structure tree")
        pdf.close()
        return
    report.ok("PDF Tags", "Logical structure tree (StructTreeRoot) present")

    headings = []
    figures = []
    tables = []

    def visit(tag_name, elem):
        if re.fullmatch(r"H[1-6]", tag_name):
            headings.append(int(tag_name[1]))
        elif tag_name == "Figure":
            alt = elem.get("/Alt")
            figures.append(str(alt) if alt else None)
        elif tag_name == "Table":
            tables.append(elem)

    kids = struct_root.get("/K")
    if kids is not None:
        _walk_struct(kids, visit)

    if not headings:
        report.warn("PDF Tags", "No heading tags (H1-H6) found in PDF structure tree")
    else:
        report.ok("PDF Tags", f"Found {len(headings)} heading tag(s) in PDF structure tree")
        if headings[0] != 1:
            report.warn("PDF Tags", f"Structure tree does not start with H1 (starts at H{headings[0]})")
        for i in range(1, len(headings)):
            if headings[i] > headings[i - 1] + 1:
                report.error("PDF Tags", f"Skipped heading level in PDF tags: H{headings[i-1]} -> H{headings[i]}")

    if figures:
        missing_alt = sum(1 for a in figures if not a)
        if missing_alt:
            report.error("PDF Tags", f"{missing_alt}/{len(figures)} Figure tag(s) missing /Alt text in the actual PDF")
        else:
            report.ok("PDF Tags", f"All {len(figures)} Figure tag(s) have /Alt text in the PDF")
    else:
        report.ok("PDF Tags", "No Figure tags found in PDF structure tree")

    if tables:
        report.ok("PDF Tags", f"Found {len(tables)} Table tag(s) in PDF structure tree")
        for idx, table in enumerate(tables, 1):
            counts = {"th": 0, "scoped": 0}

            def visit_table_cells(tag_name, elem):
                if tag_name == "TH":
                    counts["th"] += 1
                    if _struct_attr(elem, "/Scope") is not None:
                        counts["scoped"] += 1

            table_kids = table.get("/K")
            if table_kids is not None:
                _walk_struct(table_kids, visit_table_cells)

            if counts["th"] == 0:
                report.error("PDF Tags", f"Table {idx}: no TH (header cell) tags found in PDF structure")
            elif counts["scoped"] < counts["th"]:
                report.warn("PDF Tags", f"Table {idx}: {counts['th']-counts['scoped']}/{counts['th']} TH cell(s) missing /Scope attribute")
            else:
                report.ok("PDF Tags", f"Table {idx}: all TH cells have /Scope set")
    else:
        report.ok("PDF Tags", "No Table tags found in PDF structure tree")

    total_widgets = 0
    untagged_widgets = 0
    missing_tu = 0
    for page in pdf.pages:
        annots = page.get("/Annots")
        if not annots:
            continue
        for annot in annots:
            if str(annot.get("/Subtype", "")) != "/Widget":
                continue
            total_widgets += 1
            if "/StructParent" not in annot:
                untagged_widgets += 1
            if not annot.get("/TU"):
                missing_tu += 1

    if total_widgets:
        if untagged_widgets:
            report.error("PDF Tags", f"{untagged_widgets}/{total_widgets} form field(s) not tagged (missing /StructParent)")
        else:
            report.ok("PDF Tags", f"All {total_widgets} form field(s) tagged")
        if missing_tu:
            report.warn("PDF Tags", f"{missing_tu}/{total_widgets} form field(s) missing tooltip /TU description")
    else:
        report.ok("PDF Tags", "No form fields found")

    pdf.close()


# ---------------------------------------------------------------------------
# Layer 3: VeraPDF (PDF/UA-1)
# ---------------------------------------------------------------------------

def find_verapdf() -> str | None:
    """Locate the verapdf binary."""
    # Common install locations
    candidates = [
        "verapdf",
        "/opt/verapdf/verapdf",
        "/usr/local/bin/verapdf",
        os.path.expanduser("~/verapdf/verapdf"),
    ]
    for c in candidates:
        try:
            result = subprocess.run([c, "--version"], capture_output=True, timeout=5)
            if result.returncode == 0:
                return c
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
    return None


# ---------------------------------------------------------------------------
# Layer 3: Visual similarity check (source PDF vs output PDF)
# ---------------------------------------------------------------------------

def check_visual_similarity(source_pdf: str, output_pdf: str, report: CheckReport):
    """
    Render both PDFs to images and compare them using pixel-level diff.
    Uses Pillow only — no API, no internet, completely free.

    Threshold: >85% pixel similarity per page = PASS.
    """
    try:
        from PIL import Image, ImageChops
        import math
    except ImportError:
        report.warn("Visual Match", "Pillow not installed — skipping visual check (pip install pillow)")
        return

    try:
        src_images = convert_from_path(source_pdf, dpi=72)   # low DPI for speed
        out_images = convert_from_path(output_pdf, dpi=72)
    except Exception as e:
        report.warn("Visual Match", f"Could not render PDFs for comparison: {e}")
        return

    page_count = min(len(src_images), len(out_images))
    if len(src_images) != len(out_images):
        report.warn("Visual Match", f"Page count mismatch: source={len(src_images)}, output={len(out_images)}")

    total_similarity = 0.0
    low_similarity_pages = []

    for i in range(page_count):
        src = src_images[i].convert("L")   # grayscale
        out = out_images[i].convert("L")

        # Resize output to match source dimensions
        if src.size != out.size:
            out = out.resize(src.size, Image.LANCZOS)

        diff = ImageChops.difference(src, out)
        total_pixels = diff.width * diff.height

        # Count pixels within acceptable tolerance (diff < 30/255)
        hist = diff.histogram()  # 256-bucket histogram for grayscale
        similar_pixels = sum(hist[:30])
        similarity = similar_pixels / total_pixels
        total_similarity += similarity

        if similarity < 0.85:
            low_similarity_pages.append((i + 1, similarity))

    avg_similarity = total_similarity / page_count if page_count else 0
    pct = avg_similarity * 100

    if not low_similarity_pages:
        report.ok("Visual Match", f"All {page_count} pages visually similar to source (avg {pct:.1f}% match)")
    else:
        for page_num, sim in low_similarity_pages:
            report.warn("Visual Match", f"Page {page_num}: {sim*100:.1f}% visual match — layout may differ from original")
        if avg_similarity < 0.70:
            report.error("Visual Match", f"Overall visual match too low: {pct:.1f}% — output looks significantly different from source")
        else:
            report.warn("Visual Match", f"Average visual match: {pct:.1f}% — some pages differ from source")


def run_verapdf(pdf_path: str, report: CheckReport):
    """Run VeraPDF PDF/UA-1 validation. Skips gracefully if not installed."""
    verapdf = find_verapdf()
    if not verapdf:
        report.warn(
            "PDF/UA (VeraPDF)",
            "VeraPDF not found — skipping PDF/UA-1 validation.\n"
            "    Install: sudo apt install default-jre && "
            "wget https://docs.verapdf.org/install/ (see README)"
        )
        return

    print("Running VeraPDF PDF/UA-1 validation...")
    with tempfile.NamedTemporaryFile(suffix=".json", delete=False) as f:
        out_file = f.name

    try:
        result = subprocess.run(
            [verapdf, "--flavour", "ua1", "--format", "json", "--output", out_file, pdf_path],
            capture_output=True,
            text=True,
            timeout=120,
        )

        if not os.path.exists(out_file):
            report.warn("PDF/UA (VeraPDF)", f"VeraPDF produced no output. stderr: {result.stderr[:200]}")
            return

        with open(out_file) as f:
            data = json.load(f)

        jobs = data.get("report", {}).get("jobs", [])
        if not jobs:
            report.warn("PDF/UA (VeraPDF)", "No jobs in VeraPDF output")
            return

        job = jobs[0]
        validation = job.get("validationResult", {})
        compliant = validation.get("compliant", False)
        passed_rules = validation.get("passedRules", 0)
        failed_rules = validation.get("failedRules", 0)
        passed_checks = validation.get("passedChecks", 0)
        failed_checks = validation.get("failedChecks", 0)

        if compliant:
            report.ok("PDF/UA (VeraPDF)", f"PDF/UA-1 COMPLIANT — {passed_rules} rules passed, {passed_checks} checks")
        else:
            report.error("PDF/UA (VeraPDF)", f"PDF/UA-1 FAILED — {failed_rules} rules failed, {failed_checks} checks failed")

            # Report specific failures (top 10)
            rule_violations = validation.get("details", {}).get("ruleSummaries", [])
            for rule in rule_violations[:10]:
                clause = rule.get("clause", "?")
                test_num = rule.get("testNumber", "?")
                desc = rule.get("description", "")
                fails = rule.get("failedChecks", 0)
                report.error("PDF/UA (VeraPDF)", f"  Clause {clause}.{test_num} ({fails} failures): {desc}")

    except subprocess.TimeoutExpired:
        report.warn("PDF/UA (VeraPDF)", "VeraPDF timed out")
    except Exception as e:
        report.warn("PDF/UA (VeraPDF)", f"VeraPDF error: {e}")
    finally:
        if os.path.exists(out_file):
            os.unlink(out_file)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    if len(sys.argv) < 3:
        print("Usage: python3 ada_check.py source.pdf remediated.docx [remediated.pdf]")
        sys.exit(1)

    source_pdf = sys.argv[1]
    docx_path = sys.argv[2]
    pdf_path = sys.argv[3] if len(sys.argv) > 3 else None

    for path in [source_pdf, docx_path]:
        if not os.path.exists(path):
            print(f"Error: file not found: {path}")
            sys.exit(1)

    report = CheckReport(source_pdf=source_pdf, docx_path=docx_path, pdf_path=pdf_path)

    # Layer 1 — always runs
    run_docx_checks(source_pdf, docx_path, report)

    # Layer 2, 3 & 4 — only if an output PDF was provided
    if pdf_path:
        if not os.path.exists(pdf_path):
            report.warn("PDF Tags", f"PDF not found: {pdf_path} — skipping PDF tag, PDF/UA, and visual checks")
        else:
            check_pdf_tags(pdf_path, report)
            run_verapdf(pdf_path, report)
            check_visual_similarity(source_pdf, pdf_path, report)
    else:
        report.warn(
            "PDF Tags",
            "No output PDF provided — export from Word and re-run:\n"
            "    python3 ada_check.py source.pdf output.docx output.pdf"
        )

    passed = report.print_report()
    sys.exit(0 if passed else 1)


if __name__ == "__main__":
    main()
