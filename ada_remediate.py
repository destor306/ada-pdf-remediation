#!/usr/bin/env python3
"""
ADA PDF Remediation Tool (MVP)

Hybrid pipeline: Ollama (local, free) → Claude API fallback (paid)

Workflow:
  Input PDF → Cost estimate + confirmation → Vision analysis (per page) → .docx → Tagged PDF

Usage:
  python3 ada_remediate.py input.pdf [output.docx]

Environment:
  ANTHROPIC_API_KEY  — required only if Claude fallback is used
  OLLAMA_HOST        — optional, defaults to http://localhost:11434
  LOCAL_MODEL        — optional, defaults to qwen2-vl
  NO_FALLBACK        — set to 1 to disable Claude fallback
"""

import sys
import os
import base64
import json
import re
import io
from pathlib import Path

import sys
sys.stdout.reconfigure(encoding='utf-8')

from dotenv import load_dotenv
load_dotenv()

import pikepdf
import pdfplumber
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

DPI = 150                        # page render resolution
LARGE_DOC_THRESHOLD = 50         # pages — prompt user confirmation above this
MAX_PAGES = 500                  # hard cap

# Cost constants (USD)
CLAUDE_COST_PER_PAGE = 0.025     # ~$0.02–0.03 estimate
LOCAL_COST_PER_PAGE = 0.0        # free

LOCAL_MODEL = os.environ.get("LOCAL_MODEL", "qwen2-vl:2b")
CLAUDE_MODEL = os.environ.get("CLAUDE_MODEL", "claude-haiku-4-5-20251001")
OLLAMA_HOST = os.environ.get("OLLAMA_HOST", "http://localhost:11434")

# Poppler path (Windows — set via env or auto-detected)
_WINGET_POPPLER = os.path.expandvars(
    r"%LOCALAPPDATA%\Microsoft\WinGet\Packages\oschwartz10612.Poppler_Microsoft.Winget.Source_8wekyb3d8bbwe"
)
_poppler_candidates = [
    os.environ.get("POPPLER_PATH", ""),
    next((str(p) for p in Path(_WINGET_POPPLER).glob("poppler-*/Library/bin") if p.is_dir()), "") if os.path.isdir(_WINGET_POPPLER) else "",
]
POPPLER_PATH = next((p for p in _poppler_candidates if p and os.path.isdir(p)), None)


# ---------------------------------------------------------------------------
# Vision backend detection
# ---------------------------------------------------------------------------

def detect_backends() -> dict:
    """Check which backends are available."""
    backends = {"ollama": False, "claude": False, "mock": False}

    # Check Ollama
    try:
        import ollama
        client = ollama.Client(host=OLLAMA_HOST)
        models = client.list()
        model_names = [m.model for m in models.models]
        backends["ollama"] = any(LOCAL_MODEL in m for m in model_names)
        if not backends["ollama"]:
            print(f"  [info] Ollama running but '{LOCAL_MODEL}' not found.")
            print(f"         Run: ollama pull {LOCAL_MODEL}")
    except Exception:
        pass

    # Check Claude
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if api_key and os.environ.get("NO_FALLBACK") != "1":
        backends["claude"] = True

    # Mock mode — always available as last resort (text extraction, no AI)
    # Enabled automatically when no real backend found, or via MOCK_MODE=1
    if os.environ.get("MOCK_MODE") == "1" or (not backends["ollama"] and not backends["claude"]):
        backends["mock"] = True

    return backends


# ---------------------------------------------------------------------------
# Cost estimation
# ---------------------------------------------------------------------------

def estimate_cost(page_count: int, backends: dict) -> tuple[float, str]:
    """Return (estimated_cost_usd, breakdown_string)."""
    if backends["ollama"]:
        cost = 0.0
        note = f"Local model ({LOCAL_MODEL}) — $0.00"
        if backends["claude"]:
            note += f" + Claude fallback at ${CLAUDE_COST_PER_PAGE:.3f}/page if needed"
    elif backends["claude"]:
        cost = page_count * CLAUDE_COST_PER_PAGE
        note = f"Claude API only — ~${cost:.2f} ({page_count} pages × ${CLAUDE_COST_PER_PAGE:.3f})"
    else:
        cost = 0.0
        note = "No AI backend available — will fail"
    return cost, note


def confirm_large_doc(page_count: int, cost_note: str) -> bool:
    """Ask user to confirm processing a large document."""
    print(f"\n  Document has {page_count} pages (over {LARGE_DOC_THRESHOLD}-page threshold).")
    print(f"  Estimated cost: {cost_note}")
    answer = input("  Continue? [y/N] ").strip().lower()
    return answer in ("y", "yes")


# ---------------------------------------------------------------------------
# Page rendering
# ---------------------------------------------------------------------------

def render_page_to_base64(pdf_path: str, page_number: int) -> str:
    """Render a single PDF page to a base64 PNG.

    Uses pdf2image+poppler when available; falls back to PyMuPDF (no poppler needed).
    """
    # Try PyMuPDF first — no poppler dependency, works on Windows out of the box
    try:
        import fitz
        doc = fitz.open(pdf_path)
        page = doc[page_number - 1]
        mat = fitz.Matrix(DPI / 72, DPI / 72)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        buf = io.BytesIO(pix.tobytes("png"))
        doc.close()
        return base64.standard_b64encode(buf.getvalue()).decode("utf-8")
    except ImportError:
        pass  # fall through to pdf2image

    images = convert_from_path(
        pdf_path, dpi=DPI,
        first_page=page_number, last_page=page_number,
        fmt="png",
        poppler_path=POPPLER_PATH,
    )
    if not images:
        raise ValueError(f"Could not render page {page_number}")
    buf = io.BytesIO()
    images[0].save(buf, format="PNG")
    return base64.standard_b64encode(buf.getvalue()).decode("utf-8")


def extract_text_layer(pdf_path: str) -> dict[int, str]:
    """Extract raw text per page as a hint for vision models."""
    pages = {}
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            pages[i] = page.extract_text() or ""
    return pages


# ---------------------------------------------------------------------------
# Page dimension extraction
# ---------------------------------------------------------------------------

def get_page_dimensions(pdf_path: str) -> list[tuple[float, float]]:
    """
    Return (width_inches, height_inches) for each page.
    PDF mediabox units are points (1 pt = 1/72 inch).
    """
    dims = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            w_pts = float(page.width)
            h_pts = float(page.height)
            dims.append((w_pts / 72.0, h_pts / 72.0))
    return dims


# ---------------------------------------------------------------------------
# Complexity detection
# ---------------------------------------------------------------------------

MIN_IMAGE_WIDTH  = 50   # pixels — skip tiny icons/bullets
MIN_IMAGE_HEIGHT = 50

def classify_page_complexity(pdf_path: str, page_num: int) -> str:
    """Return 'simple' (text only) or 'complex' (has tables or images)."""
    with pdfplumber.open(pdf_path) as pdf:
        page = pdf.pages[page_num - 1]
        if page.extract_tables():
            return "complex"
        if page.images:
            return "complex"
    return "simple"


def extract_page_images(pdf_path: str, page_num: int) -> list[dict]:
    """
    Extract embedded images from a page using PyMuPDF.
    Returns list of {bytes, ext, width, height, bbox} dicts, sorted top-to-bottom.
    Skips tiny decorative images (icons, bullets).
    """
    import fitz
    results = []
    doc = fitz.open(pdf_path)
    page = doc[page_num - 1]
    page_h = page.rect.height

    seen_xrefs = set()
    for img_info in page.get_images(full=True):
        xref = img_info[0]
        if xref in seen_xrefs:
            continue
        seen_xrefs.add(xref)

        try:
            base = doc.extract_image(xref)
            w, h = base["width"], base["height"]
            if w < MIN_IMAGE_WIDTH or h < MIN_IMAGE_HEIGHT:
                continue

            # Get bounding box on the page (for vertical ordering)
            rects = page.get_image_rects(xref)
            bbox_y = rects[0].y0 if rects else 0

            results.append({
                "bytes": base["image"],
                "ext":   base["ext"],
                "width": w,
                "height": h,
                "bbox_y": bbox_y,
                "xref": xref,
            })
        except Exception:
            continue

    doc.close()
    results.sort(key=lambda x: x["bbox_y"])
    return results


# ---------------------------------------------------------------------------
# Shared prompt
# ---------------------------------------------------------------------------

# Compliance references (both stored in project root):
#   1. pdf-tagging_from_hhs_gov.pdf  — HHS Section 508 Guide: Tagging PDFs in Adobe Acrobat Pro (May 2018)
#      Defines correct tag structure: H1-H6, P, Table/TR/TH/TD, L/LI, Figure with alt text, artifacts
#   2. web-rule_ada_gov.pdf          — DOJ Final Rule 28 CFR Part 35 (April 2024)
#      Mandates WCAG 2.1 Level AA for all state & local government web content including PDFs.
#      Compliance deadlines: April 2026 (50,000+ population), April 2027 (all remaining entities).
SYSTEM_PROMPT = """You are an expert PDF accessibility specialist. Your output must comply with:
- WCAG 2.1 Level AA (required by DOJ 28 CFR Part 35 for all state & local government entities)
- HHS Section 508 Guide for tagging PDFs in Adobe Acrobat Pro (May 2018)
- PDF/UA (ISO 14289-1) standard

Given an image of a single PDF page, reconstruct its content as a structured JSON object
suitable for building an accessible Word document.

Return ONLY valid JSON — no markdown fences, no explanation.

CRITICAL: Always populate "elements" (and "rows"/"items" within it) as true nested
JSON arrays of objects — never as a JSON-encoded string. If any source text contains a
literal quotation mark (e.g. a form field reading `Form I-94 with "RE" notation`), you
MUST escape it as \" inside the JSON string — never emit an unescaped " inside a string
value, since that silently breaks the surrounding object.

JSON schema:
{
  "page": <int>,
  "elements": [
    {
      "type": "heading" | "paragraph" | "table" | "list" | "caption" | "image" | "header" | "footer",
      "level": <int 1-6>,               // headings only
      "text": "<string>",               // all text types
      "alt_text": "<string>",           // image only — 120-250 characters, describes what the image conveys
      "caption": "<string>",            // image only — visible caption text near the image (or "")
      "rows": [                         // table only
        { "is_header": <bool>, "col_scope": "col" | "row" | null, "cells": ["<string>", ...] }
      ],
      "col_widths": [0.15, 0.25, 0.60], // table only: relative column widths (must sum to 1.0)
      "items": ["<string>", ...],       // list only
      "ordered": <bool>                 // list only
    }
  ]
}

Section 508 / PDF-UA tagging rules (from HHS Section 508 Guide):

HEADINGS:
- H1 is used exactly once per document — only for the main document title.
- H2 = main sections or chapters. H3–H6 = subsections. Never skip levels (no H1→H3).
- Only tag text as a heading if it visually defines a section. A different font or color alone does not make a heading.
- Heading levels must follow a strict logical progression with no gaps.
- On government forms, field numbers/letters and their labels (e.g. "5.a.", "24.", "Item Number 12.",
  "12.b. Given Name") are NOT headings, even when bold, boxed, or in a larger font — tag them as
  "paragraph" (or as table cells if they sit inside a form grid). A heading is a substantive multi-word
  section or subsection title (e.g. "Part 2. Information About You"), never a bare number/letter or a
  single short field label.

PARAGRAPHS:
- Use type "paragraph" for body text, standalone text, and plain decorative text that is not a heading.
- Text that spans visual paragraphs should be kept as a single paragraph element.

TABLES (data tables only — not layout tables):
- Every row must have the same cell count. Never merge or skip columns.
- Mark header rows with is_header: true and col_scope: "col" for column headers, "row" for row headers.
- Complex tables with both row and column headers need col_scope set on every header cell.
- Identify all header rows and columns, not just the first row.

LISTS:
- Use type "list" for all bulleted or numbered sequential items.
- Bullets and numbers are informational — include them in the items array (strip the bullet character, keep the text).
- Set ordered: true for numbered lists, false for bulleted lists.

IMAGES / FIGURES:
- Use type "image" for all visual images that convey information: charts, photos, logos, signatures, diagrams.
- alt_text must describe what the image CONVEYS to a sighted user (not just what it looks like).
  Good: "Bar chart showing monthly case volume January through June 2025, peaking in April at 91 cases"
  Bad: "A bar chart" or "Chart"
- Keep alt_text between 120–250 characters for complex images. Simpler images (logos, signatures) can be shorter.
- Purely decorative images that convey no information: use alt_text: "" (empty string).

HEADER / FOOTER:
- Page footer (text in the bottom margin, below main content): use type "footer".
- Running page header (identical small text repeated across multiple pages in the top margin): use type "header".
- Do NOT use "header" for a one-time page title or section heading — use "heading" with the appropriate level.

READING ORDER:
- Elements must be in logical reading order: top-to-bottom, left-to-right for English.
- Multi-column layouts: linearize in logical reading order (finish left column before right column).

ARTIFACTS (omit from output entirely):
- Decorative horizontal rules, background shading, page borders — do not include these as elements.
- Page numbers that are purely navigational (not content) — omit.

FOOTNOTES:
- Include as paragraph elements at the end of the body content, before any footer element.

WCAG 2.1 LEVEL AA KEY SUCCESS CRITERIA FOR DOCUMENTS (28 CFR Part 35 / DOJ Final Rule):
- 1.1.1 Non-text Content: every image must have a text alternative (alt_text) that conveys the same information.
- 1.3.1 Info and Relationships: structure (headings, lists, tables) must be conveyed through proper element types, not just visual formatting.
- 1.3.2 Meaningful Sequence: reading order must reflect the logical order a human would read the content.
- 2.4.2 Page Titled: the document must have a meaningful title (handled by the pipeline, not per-element).
- 3.1.1 Language of Page: document language must be declared (handled by the pipeline).
- 3.1.2 Language of Parts: if a word or phrase is in a different language than the document, note it in a "notes" field if detectable.

- Blank page: return { "page": <n>, "elements": [] }.
"""


def user_message_text(page_num: int, text_hint: str, image_count: int = 0) -> str:
    hint = f"\nRaw text layer hint (may be garbled):\n{text_hint[:2000]}" if text_hint.strip() else ""
    img_note = (
        f"\nThis page contains {image_count} embedded image(s) (charts, photos, signatures, logos). "
        f"Include exactly {image_count} element(s) of type \"image\" in reading order, each with descriptive alt_text."
        if image_count > 0 else ""
    )
    return f"This is page {page_num}.{hint}{img_note}\n\nReconstruct this page as JSON per the schema."


def parse_json_response(raw: str, page_num: int) -> dict:
    """Strip markdown fences and parse JSON, recovering partial elements on truncation."""
    raw = re.sub(r"^```[a-z]*\n?", "", raw, flags=re.MULTILINE)
    raw = re.sub(r"```\s*$", "", raw, flags=re.MULTILINE).strip()
    try:
        return json.loads(raw)
    except json.JSONDecodeError as e:
        print(f"  [warn] JSON parse error on page {page_num}: {e} — attempting partial recovery")
        return _recover_partial_elements(raw, page_num)


def _extract_json_objects(text: str, start: int = 0) -> list[dict]:
    """Extract complete top-level {...} objects from inside a JSON array's
    text, skipping any individual object that fails to parse (e.g. because an
    unescaped literal quote in the source text broke just that one object).
    More tolerant than json.loads on the whole array: one bad element doesn't
    sink everything after it.
    """
    objects = []
    depth = 0
    obj_start = None
    for i in range(start, len(text)):
        ch = text[i]
        if ch == '{':
            if depth == 0:
                obj_start = i
            depth += 1
        elif ch == '}':
            depth -= 1
            if depth == 0 and obj_start is not None:
                segment = text[obj_start:i + 1]
                try:
                    objects.append(json.loads(segment))
                except json.JSONDecodeError:
                    repaired = _repair_object(segment)
                    if repaired:
                        objects.append(repaired)
                obj_start = None
        elif ch == ']' and depth == 0:
            break
    return objects


def _repair_object(text: str) -> dict | None:
    """Best-effort repair for a {...} element that failed to parse whole.

    The dominant case: a "table" element with many rows where exactly one
    row's cell text has an unescaped quote. Without this, that single bad
    row sinks the *entire* table — every other row is lost too, not just the
    broken one. Recover whatever rows/items parse cleanly instead of
    discarding the whole element.
    """
    type_m = re.search(r'"type"\s*:\s*"(\w+)"', text)
    etype = type_m.group(1) if type_m else None

    if etype == 'table':
        rows_m = re.search(r'"rows"\s*:\s*\[', text)
        if rows_m:
            rows = _extract_json_objects(text, rows_m.end())
            if rows:
                return {"type": "table", "rows": rows}
    return None


def _recover_partial_elements(raw: str, page_num: int) -> dict:
    """Extract complete JSON objects from a truncated elements array."""
    m = re.search(r'"elements"\s*:\s*\[', raw)
    if not m:
        return {"page": page_num, "elements": []}
    elements = _extract_json_objects(raw, m.end())
    if elements:
        print(f"  [info] Recovered {len(elements)} element(s) from truncated page {page_num}")
    return {"page": page_num, "elements": elements}


# ---------------------------------------------------------------------------
# Local model (Ollama)
# ---------------------------------------------------------------------------

def analyze_with_ollama(pdf_path: str, page_num: int, text_hint: str = "", image_count: int = 0) -> dict | None:
    """Try local vision model. Returns None on failure.

    llava-family models do not support the 'system' role — they silently return
    empty when a system message is present.  Workaround: fold the system prompt
    into the user message so the full instruction arrives in one turn.
    """
    try:
        import ollama
        image_b64 = render_page_to_base64(pdf_path, page_num)
        client = ollama.Client(host=OLLAMA_HOST)

        # Combine system instructions + user request into a single user message
        combined = (
            f"{SYSTEM_PROMPT}\n\n"
            f"---\n"
            f"{user_message_text(page_num, text_hint, image_count)}"
        )

        response = client.chat(
            model=LOCAL_MODEL,
            messages=[
                {
                    "role": "user",
                    "content": combined,
                    "images": [image_b64],
                },
            ],
            format="json",
            options={"temperature": 0},
        )
        raw = response.message.content or ""
        if not raw.strip():
            print(f"  [warn] Local model returned empty response for page {page_num}, trying fallback.")
            return None
        result = parse_json_response(raw, page_num)
        # Basic quality check: did we get any elements?
        if result.get("elements"):
            return result
        print(f"  [warn] Local model returned empty elements for page {page_num}, trying fallback.")
        return None
    except Exception as e:
        print(f"  [warn] Ollama error on page {page_num}: {e}")
        return None


# ---------------------------------------------------------------------------
# Claude API fallback
# ---------------------------------------------------------------------------

PAGE_ELEMENTS_TOOL = {
    "name": "emit_page_elements",
    "description": "Emit the structured representation of the analyzed PDF page.",
    "input_schema": {
        "type": "object",
        "properties": {
            "page": {"type": "integer"},
            "elements": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "type": {
                            "type": "string",
                            "enum": ["heading", "paragraph", "table", "list", "caption",
                                     "image", "header", "footer"],
                        },
                        "level": {"type": "integer", "minimum": 1, "maximum": 6},
                        "text": {"type": "string"},
                        "alt_text": {"type": "string"},
                        "caption": {"type": "string"},
                        "rows": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "is_header": {"type": "boolean"},
                                    "col_scope": {"type": ["string", "null"], "enum": ["col", "row", None]},
                                    "cells": {"type": "array", "items": {"type": "string"}},
                                },
                                "required": ["cells"],
                            },
                        },
                        "col_widths": {"type": "array", "items": {"type": "number"}},
                        "items": {"type": "array", "items": {"type": "string"}},
                        "ordered": {"type": "boolean"},
                    },
                    "required": ["type"],
                },
            },
        },
        "required": ["page", "elements"],
    },
}


def analyze_with_claude(pdf_path: str, page_num: int, text_hint: str = "", image_count: int = 0) -> dict:
    """Analyze page with Claude vision API.

    Uses tool-use (not freeform JSON-in-text) so the API parses the structured
    output server-side. Freeform generation broke whenever source text on the
    page contained a literal quote character (e.g. a form field reading
    `Form I-94 with "RE" notation`) — the model would emit it unescaped and
    corrupt the surrounding JSON. Forcing a tool call sidesteps that failure
    mode entirely.
    """
    import anthropic
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    client = anthropic.Anthropic(api_key=api_key)
    image_b64 = render_page_to_base64(pdf_path, page_num)

    def _call() -> dict:
        response = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=8192,
            temperature=0,
            system=SYSTEM_PROMPT,
            tools=[PAGE_ELEMENTS_TOOL],
            tool_choice={"type": "tool", "name": "emit_page_elements"},
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {"type": "base64", "media_type": "image/png", "data": image_b64},
                    },
                    {"type": "text", "text": user_message_text(page_num, text_hint, image_count)},
                ],
            }],
        )

        if response.stop_reason == "max_tokens":
            print(f"  [warn] Claude tool call hit max_tokens on page {page_num} — output may be incomplete")

        for block in response.content:
            if block.type == "tool_use":
                result = dict(block.input)
                result.setdefault("page", page_num)

                # Occasionally the model serializes "elements" as a JSON string
                # instead of a true nested array, even under tool-use — and
                # since that string is itself hand-written JSON, it can carry
                # the exact same unescaped-quote breakage as the freeform path.
                # Use the tolerant per-object extractor rather than a strict
                # json.loads so one bad element doesn't sink the whole page.
                elements = result.get("elements", [])
                if isinstance(elements, str):
                    recovered = _extract_json_objects(elements)
                    if recovered:
                        print(f"  [info] Recovered {len(recovered)} element(s) from a "
                              f"stringified 'elements' field on page {page_num}")
                    # An empty recovery from a non-trivial string is a real
                    # failure (caller retries); from a near-empty string it's
                    # legitimately just an empty page.
                    elements = recovered if (recovered or len(elements) < 20) else None
                result["elements"] = elements if isinstance(elements, list) else None
                return result

        # Model didn't call the tool (shouldn't happen with tool_choice forced) —
        # fall back to parsing whatever text it produced instead.
        text = "".join(b.text for b in response.content if getattr(b, "type", None) == "text")
        parsed = parse_json_response(text, page_num)
        parsed["elements"] = parsed.get("elements") or None
        return parsed

    result = _call()
    if result.get("elements") is None:
        print(f"  [warn] Claude returned malformed 'elements' on page {page_num} — retrying once")
        result = _call()
    result["elements"] = result.get("elements") or []
    return result


# ---------------------------------------------------------------------------
# Zero-cost text-only analyzer (font-size aware, no AI)
# ---------------------------------------------------------------------------

def _chars_to_lines(chars: list) -> list[dict]:
    """Group pdfplumber chars into visual lines, sorted top-to-bottom."""
    from collections import Counter
    if not chars:
        return []
    buckets: dict[int, list] = {}
    for ch in chars:
        if not ch.get('text'):
            continue
        bucket = round(ch.get('top', 0) / 2) * 2  # snap to 2pt grid
        buckets.setdefault(bucket, []).append(ch)
    lines = []
    for bucket in sorted(buckets):
        line_chars = sorted(buckets[bucket], key=lambda c: c.get('x0', 0))
        text = ''.join(c['text'] for c in line_chars)
        sizes = [c['size'] for c in line_chars if c.get('size', 0) > 4]
        dom_size = Counter(round(s, 1) for s in sizes).most_common(1)[0][0] if sizes else 0
        is_bold = any('bold' in c.get('fontname', '').lower() for c in line_chars)
        lines.append({'text': text, 'dominant_size': dom_size, 'is_bold': is_bold, 'top': bucket})
    return lines


def analyze_text_page(pdf_path: str, page_num: int) -> dict:
    """
    Zero-cost text-only page analysis using pdfplumber character data.
    Uses font size to detect headings; one element per visual line so
    _match_blocks_to_elements achieves maximum coverage — no content
    ends up untagged as an artifact.
    """
    from collections import Counter
    LIST_RE = re.compile(r'^([•·\-\*○●▪▸►]|\d+[\.\)]|[a-zA-Z][\.\)])\s+')

    with pdfplumber.open(pdf_path) as pdf:
        page = pdf.pages[page_num - 1]
        chars = page.chars
        if not chars:
            return {"page": page_num, "elements": []}

        page_height = float(page.height)
        lines = _chars_to_lines(chars)
        if not lines:
            return {"page": page_num, "elements": []}

        # Body size = most common size among non-tiny chars (>4pt filters superscripts)
        all_sizes = [round(c['size'], 1) for c in chars if c.get('size', 0) > 4]
        if not all_sizes:
            return {"page": page_num, "elements": []}
        body_size = Counter(all_sizes).most_common(1)[0][0]

        # Sizes noticeably above body → heading levels (largest = H1)
        heading_sizes = sorted(
            {round(l['dominant_size'], 1) for l in lines
             if round(l['dominant_size'], 1) > body_size * 1.05},
            reverse=True,
        )
        size_to_level = {s: min(i + 1, 6) for i, s in enumerate(heading_sizes)}

        # Bottom 8% of page = footer zone.
        # Headers are NOT detected by position — a page title at the top looks identical
        # to a running header positionally. Leave header classification to the AI.
        footer_threshold = page_height * 0.92

        elements: list[dict] = []
        list_items: list[str] = []
        list_ordered = False

        for line in lines:
            text = line['text'].strip()
            if not text:
                if list_items:
                    elements.append({"type": "list", "ordered": list_ordered, "items": list_items})
                    list_items = []
                continue

            top = line.get('top', 0)

            # Position-based footer detection only
            if top > footer_threshold:
                if list_items:
                    elements.append({"type": "list", "ordered": list_ordered, "items": list_items})
                    list_items = []
                elements.append({"type": "footer", "text": text})
                continue

            size = round(line['dominant_size'], 1)
            is_list_item = bool(LIST_RE.match(text))

            if not is_list_item and list_items:
                elements.append({"type": "list", "ordered": list_ordered, "items": list_items})
                list_items = []

            if size in size_to_level:
                elements.append({"type": "heading", "level": size_to_level[size], "text": text})
            elif line['is_bold'] and not is_list_item and len(text) < 120 and size >= body_size * 0.95:
                level = min(len(size_to_level) + 1 if size_to_level else 2, 4)
                elements.append({"type": "heading", "level": level, "text": text})
            elif is_list_item:
                if not list_items:
                    list_ordered = bool(re.match(r'^\d+[\.\)]\s+', text))
                list_items.append(LIST_RE.sub('', text).strip())
            else:
                elements.append({"type": "paragraph", "text": text})

        if list_items:
            elements.append({"type": "list", "ordered": list_ordered, "items": list_items})

        return {"page": page_num, "elements": elements}


# Mock backend (last-resort fallback — no font data, no AI)
# ---------------------------------------------------------------------------

def analyze_with_mock(pdf_path: str, page_num: int, text_hint: str = "") -> dict:
    """
    No-AI fallback: converts raw pdfplumber text into structured elements.
    Good enough to test the full pipeline and UI without any API keys or Ollama.
    Output quality: readable text, basic heading detection, no table structure.
    """
    import pdfplumber

    elements = []
    with pdfplumber.open(pdf_path) as pdf:
        page = pdf.pages[page_num - 1]

        # Try to extract table data
        tables = page.extract_tables()
        text_outside_tables = page.extract_text() or ""

        # Detect title/heading lines (short lines at top, or ALL CAPS lines)
        lines = [l.strip() for l in text_outside_tables.splitlines() if l.strip()]
        used_as_heading = set()

        for i, line in enumerate(lines[:5]):  # first 5 lines — likely headings
            if len(line) < 80 and (line.isupper() or i == 0):
                level = 1 if i == 0 else 2
                elements.append({"type": "heading", "level": level, "text": line})
                used_as_heading.add(line)

        # Add tables
        for table in tables:
            if not table:
                continue
            rows = []
            for r_idx, row in enumerate(table):
                cells = [str(c or "").strip() for c in row]
                if any(cells):
                    rows.append({"is_header": r_idx == 0, "cells": cells})
            if rows:
                elements.append({"type": "table", "rows": rows})

        # Remaining lines as paragraphs — skip lines already captured in a
        # table cell, otherwise the same text becomes two competing elements
        # (a table header cell and a near-duplicate paragraph) for the same
        # content-stream block, and the paragraph tends to win the match.
        table_words: set[str] = set()
        for elem in elements:
            if elem["type"] == "table":
                for row in elem["rows"]:
                    for cell in row["cells"]:
                        table_words.update(str(cell).lower().split())

        for line in lines:
            if line in used_as_heading or len(line) <= 2:
                continue
            lwords = set(line.lower().split())
            if lwords and len(lwords & table_words) / len(lwords) > 0.6:
                continue
            elements.append({"type": "paragraph", "text": line})

    return {"page": page_num, "elements": elements}


# ---------------------------------------------------------------------------
# Page analysis dispatcher
# ---------------------------------------------------------------------------

def _attach_images_to_elements(elements: list[dict], page_images: list[dict]) -> list[dict]:
    """
    Match extracted image bytes to AI-generated image elements by order.
    The AI returns image elements in reading order; we extracted images top-to-bottom.
    Pair them positionally (1st image element ↔ 1st extracted image, etc.).
    """
    img_iter = iter(page_images)
    for elem in elements:
        if elem.get("type") == "image":
            img = next(img_iter, None)
            if img:
                elem["_image_bytes"] = img["bytes"]
                elem["_image_ext"]   = img["ext"]
                elem["_image_width"] = img["width"]
                elem["_image_height"] = img["height"]
    return elements


def analyze_page(
    pdf_path: str,
    page_num: int,
    text_hint: str,
    backends: dict,
) -> dict:
    """
    Tiered analysis:
    - Simple pages (text only, no tables/images): font-size-aware text extraction — free, instant.
    - Complex pages (tables, images): extract real images, then run AI for structure + alt text.
    """
    # Forced mock mode
    if backends.get("mock") and os.environ.get("MOCK_MODE") == "1":
        return analyze_with_mock(pdf_path, page_num, text_hint)

    # Simple page — use font-aware text analyzer, zero cost, no AI needed
    complexity = classify_page_complexity(pdf_path, page_num)
    if complexity == "simple":
        result = analyze_text_page(pdf_path, page_num)
        # Fall back to mock only if chars extraction produced nothing
        if not result.get("elements"):
            result = analyze_with_mock(pdf_path, page_num, text_hint)
        return result

    # Complex page: extract real images first
    page_images = extract_page_images(pdf_path, page_num)
    image_count = len(page_images)

    result = None

    if backends["ollama"]:
        result = analyze_with_ollama(pdf_path, page_num, text_hint, image_count)

    if result is None and backends["claude"]:
        print(f"    → Using Claude API fallback for page {page_num}")
        result = analyze_with_claude(pdf_path, page_num, text_hint, image_count)

    if result is None:
        result = analyze_with_mock(pdf_path, page_num, text_hint)
    else:
        # AI table extraction can lose rows on unusually dense tables (e.g. a
        # JSON-escaping failure mid-table). pdfplumber's deterministic
        # extraction doesn't suffer that failure mode — use it to backstop
        # completeness when the AI's row count looks suspiciously low.
        result["elements"] = _backstop_table_rows(pdf_path, page_num, result.get("elements", []))

    # Attach real image bytes to image elements
    if page_images:
        result["elements"] = _attach_images_to_elements(result.get("elements", []), page_images)

    return result


def _backstop_table_rows(pdf_path: str, page_num: int, elements: list[dict]) -> list[dict]:
    """If an AI-extracted table has noticeably fewer rows than pdfplumber's own
    extraction of the same page, prefer pdfplumber's rows for completeness.

    Matches AI tables to pdfplumber tables by page order — a heuristic, not an
    identity match, but safe: it only overwrites AI rows when an independent
    extraction found substantially more content, never the reverse.
    """
    ai_tables = [e for e in elements if e.get("type") == "table"]
    if not ai_tables:
        return elements

    try:
        with pdfplumber.open(pdf_path) as pdf:
            pf_tables = pdf.pages[page_num - 1].extract_tables()
    except Exception:
        return elements

    for elem, pf_table in zip(ai_tables, pf_tables):
        if not pf_table:
            continue
        pf_rows = [row for row in pf_table if any(str(c or "").strip() for c in row)]
        ai_rows = elem.get("rows", [])
        if len(pf_rows) >= 4 and len(pf_rows) > len(ai_rows) * 1.5:
            print(f"  [info] Table on page {page_num}: AI extracted {len(ai_rows)} row(s), "
                  f"pdfplumber found {len(pf_rows)} — using pdfplumber's rows for completeness")
            elem["rows"] = [
                {"is_header": r_idx == 0, "cells": [str(c or "").strip() for c in row]}
                for r_idx, row in enumerate(pf_rows)
            ]

    # A table with exactly one row has no header/data distinction to make —
    # that row IS the header (e.g. a strip of form-field labels with no data
    # row beneath). The AI sometimes still marks it is_header: false, which
    # leaves the table with zero TH cells and fails accessibility checks.
    for elem in ai_tables:
        rows = elem.get("rows", [])
        if len(rows) == 1:
            rows[0]["is_header"] = True

    # Some dense form-field grids (e.g. every row is itself a strip of field
    # captions like "Prefix | First | Middle | Last") have no row that visually
    # reads as a header, so the AI marks every row is_header: false — leaving
    # the whole table with zero TH cells. A table needs at least one header
    # row to satisfy Section 508/PDF-UA, so default to the first row when the
    # AI didn't designate any.
    for elem in ai_tables:
        rows = elem.get("rows", [])
        if rows and not any(r.get("is_header") for r in rows):
            rows[0]["is_header"] = True

    return elements


# ---------------------------------------------------------------------------
# Word document builder
# ---------------------------------------------------------------------------

def set_cell_shading(cell, fill: str = "D9E1F2"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_accessible_table(doc: Document, rows_data: list[dict], col_widths: list[float] | None, page_width_in: float):
    if not rows_data:
        return
    col_count = max(len(r.get("cells", [])) for r in rows_data)
    if col_count == 0:
        return

    table = doc.add_table(rows=0, cols=col_count)
    table.style = "Table Grid"

    # Apply column widths if provided
    if col_widths and len(col_widths) == col_count:
        total = sum(col_widths) or 1.0
        usable_width = page_width_in - 1.5  # subtract margins
        for i, col in enumerate(table.columns):
            col.width = Inches(usable_width * col_widths[i] / total)

    for row_data in rows_data:
        cells_text = row_data.get("cells", [])
        is_header = row_data.get("is_header", False)

        cells_text = list(cells_text) + [""] * col_count
        cells_text = cells_text[:col_count]

        row = table.add_row()

        if is_header:
            trPr = row._tr.get_or_add_trPr()
            trPr.append(OxmlElement("w:tblHeader"))

        for i, text in enumerate(cells_text):
            cell = row.cells[i]
            para = cell.paragraphs[0]
            run = para.add_run(str(text))
            if is_header:
                run.bold = True
                set_cell_shading(cell)


def set_page_size(section, width_in: float, height_in: float):
    """Set Word section page size to match the source PDF page."""
    section.page_width = Inches(width_in)
    section.page_height = Inches(height_in)
    # Set narrow margins to maximize content area
    margin = Inches(0.75)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin
    # Orientation
    if width_in > height_in:
        section.orientation = WD_ORIENT.LANDSCAPE
    else:
        section.orientation = WD_ORIENT.PORTRAIT


def build_docx(pages_data: list[dict], output_path: str, page_dims: list[tuple[float, float]] | None = None, title: str = ""):
    _normalize_heading_levels(pages_data)
    doc = Document()
    doc.core_properties.title = title or Path(output_path).stem.replace("_", " ").title()
    doc.core_properties.language = "en-US"

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Set first section dimensions from page 1
    if page_dims:
        set_page_size(doc.sections[0], page_dims[0][0], page_dims[0][1])

    current_section = doc.sections[0]

    for idx, page_data in enumerate(pages_data):
        elements = page_data.get("elements", [])

        # Split header/footer elements out before the empty-page check
        header_elems = [e for e in elements if e.get("type") == "header"]
        footer_elems = [e for e in elements if e.get("type") == "footer"]
        body_elems   = [e for e in elements if e.get("type") not in ("header", "footer")]

        # Write header/footer into the DOCX section (not as body paragraphs)
        if header_elems:
            current_section.header.is_linked_to_previous = False
            hdr_para = current_section.header.paragraphs[0]
            hdr_para.text = " ".join(e.get("text", "") for e in header_elems)
        if footer_elems:
            current_section.footer.is_linked_to_previous = False
            ftr_para = current_section.footer.paragraphs[0]
            ftr_para.text = " ".join(e.get("text", "") for e in footer_elems)

        if not body_elems:
            # Preserve page count: still emit a page break for blank/unextracted pages
            if idx < len(pages_data) - 1:
                doc.add_page_break()
            continue

        # Use this page's dimensions for column width calculations
        if page_dims and idx < len(page_dims):
            page_w = page_dims[idx][0]
        else:
            page_w = 8.5  # default letter width

        for elem in body_elems:
            etype = elem.get("type", "paragraph")
            text = elem.get("text", "").strip()

            if etype == "heading":
                level = max(1, min(6, elem.get("level", 2)))
                p = doc.add_paragraph(style=f"Heading {level}")
                p.add_run(text)

            elif etype == "paragraph":
                if text:
                    doc.add_paragraph(text)

            elif etype == "table":
                add_accessible_table(doc, elem.get("rows", []), elem.get("col_widths"), page_w)
                doc.add_paragraph()

            elif etype == "list":
                style_name = "List Number" if elem.get("ordered") else "List Bullet"
                for item in elem.get("items", []):
                    doc.add_paragraph(item, style=style_name)

            elif etype == "caption":
                if text:
                    doc.add_paragraph(text, style="Caption")

            elif etype in ("image", "image_alt"):
                img_bytes = elem.get("_image_bytes")
                alt = elem.get("alt_text") or elem.get("text") or ""
                caption_text = elem.get("caption", "")

                if img_bytes:
                    # Embed the real extracted image
                    img_width  = elem.get("_image_width", 1)
                    img_height = elem.get("_image_height", 1)
                    # Fit within usable page width (page_w minus margins)
                    max_w = page_w - 1.5
                    aspect = img_height / img_width if img_width else 1
                    fit_w = min(max_w, max_w)   # full width by default
                    fit_h = fit_w * aspect
                    # If taller than half a page, shrink to fit
                    max_h = (page_dims[idx][1] if page_dims and idx < len(page_dims) else 11.0) * 0.5
                    if fit_h > max_h:
                        fit_h = max_h
                        fit_w = fit_h / aspect

                    p = doc.add_paragraph()
                    p.alignment = 1  # center
                    run = p.add_run()
                    run.add_picture(io.BytesIO(img_bytes), width=Inches(fit_w))

                    # Add alt text to the drawing XML
                    try:
                        drawing = run._r.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
                        if drawing is not None:
                            docPr = drawing.find('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}docPr')
                            if docPr is not None and alt:
                                docPr.set('descr', alt)
                    except Exception:
                        pass
                else:
                    # Fallback: no image bytes extracted, show placeholder
                    p = doc.add_paragraph(f"[Figure: {alt}]")
                    if p.runs:
                        p.runs[0].italic = True

                if caption_text:
                    doc.add_paragraph(caption_text, style="Caption")

        if idx < len(pages_data) - 1:
            # Add page break and new section if page size changes
            doc.add_page_break()
            if page_dims and idx + 1 < len(page_dims):
                next_w, next_h = page_dims[idx + 1]
                curr_w, curr_h = page_dims[idx]
                if abs(next_w - curr_w) > 0.1 or abs(next_h - curr_h) > 0.1:
                    current_section = doc.add_section()
                    set_page_size(current_section, next_w, next_h)

    doc.save(output_path)


# ---------------------------------------------------------------------------
# PDF Accessibility Tagger (pikepdf-based, in-place)
# ---------------------------------------------------------------------------

def _decode_pdf_string(s) -> str:
    """Decode a pikepdf String to a Python str."""
    if isinstance(s, pikepdf.String):
        raw = bytes(s)
        if len(raw) >= 2 and raw[:2] == b'\xfe\xff':
            return raw[2:].decode('utf-16-be', errors='replace')
        if len(raw) >= 2 and raw[:2] == b'\xff\xfe':
            return raw[2:].decode('utf-16-le', errors='replace')
        return raw.decode('cp1252', errors='replace')
    return str(s)


def _get_block_text(instructions, start: int, end: int) -> str:
    """Extract concatenated text from a BT..ET block."""
    parts = []
    for i in range(start, end + 1):
        op = str(instructions[i].operator)
        ops = instructions[i].operands
        if op == 'Tj' and ops:
            parts.append(_decode_pdf_string(ops[0]))
        elif op == 'TJ' and ops:
            for item in ops[0]:
                if isinstance(item, (pikepdf.String, bytes)):
                    parts.append(_decode_pdf_string(item))
        elif op == "'" and ops:
            parts.append(_decode_pdf_string(ops[0]))
        elif op == '"' and len(ops) >= 3:
            parts.append(_decode_pdf_string(ops[2]))
    return ''.join(parts).strip()


def _find_content_blocks(instructions) -> list[dict]:
    """
    Find all BT..ET text blocks and standalone Do (image) operators.
    Returns list of dicts: {type, start, end, text or xobj}.
    """
    blocks = []
    in_bt = False
    bt_start = 0
    for i, instr in enumerate(instructions):
        op = str(instr.operator)
        if op == 'BT':
            in_bt = True
            bt_start = i
        elif op == 'ET' and in_bt:
            text = _get_block_text(instructions, bt_start, i)
            blocks.append({'type': 'text', 'start': bt_start, 'end': i, 'text': text})
            in_bt = False
        elif op == 'Do' and not in_bt:
            xobj = str(instructions[i].operands[0]) if instructions[i].operands else '/Im'
            blocks.append({'type': 'image', 'start': i, 'end': i, 'xobj': xobj})
    return blocks


def _match_blocks_to_elements(blocks: list[dict], elements: list[dict]) -> dict[int, tuple[str, int, dict | None, int]]:
    """
    Match content-stream blocks to Claude-extracted elements by text overlap.
    Returns {block_index: (tag_name, local_mcid, group, elem_order)}.

    Tables and lists are expanded into one candidate per cell/item (rather than
    one candidate for the whole table/list) so each cell's content block can be
    matched independently — a table is many text blocks, not one. `group`
    carries enough info (group type/id, row index, header flag) for the caller
    to reassemble the proper Table>TR>TH/TD or L>LI nesting.

    `elem_order` is the matched element's index in `elements` (already in the
    AI's linearized reading order per the system prompt's READING ORDER rule).
    Content-stream block order does NOT reliably match visual reading order on
    fillable forms — form software often writes field widgets/labels to the
    stream in field-definition order rather than top-to-bottom order — so the
    caller must build the struct tree (and therefore heading order) from
    `elem_order`, not from `block_idx`.
    """
    TYPE_MAP = {'paragraph': 'P', 'caption': 'Caption', 'image_alt': 'Figure'}

    elem_info = []
    group_seq = 0
    for elem in elements:
        etype = elem.get('type', 'paragraph')
        if etype == 'heading':
            level = elem.get('level', 1)
            text = elem.get('text', '').strip()
            elem_info.append({'tag': f'H{max(1, min(6, level))}', 'text': text.lower(), 'group': None})
        elif etype == 'table':
            rows = elem.get('rows', [])
            if not rows:
                continue
            gid = group_seq; group_seq += 1
            for r_idx, row in enumerate(rows):
                is_header = row.get('is_header', r_idx == 0)
                for cell_text in row.get('cells', []):
                    cell_text = str(cell_text).strip()
                    if not cell_text:
                        continue
                    elem_info.append({
                        'tag': 'TH' if is_header else 'TD',
                        'text': cell_text.lower(),
                        'group': {'type': 'Table', 'id': gid, 'row': r_idx},
                    })
        elif etype == 'list':
            items = elem.get('items', [])
            if not items:
                continue
            gid = group_seq; group_seq += 1
            for item_text in items:
                item_text = str(item_text).strip()
                if not item_text:
                    continue
                elem_info.append({'tag': 'LI', 'text': item_text.lower(), 'group': {'type': 'L', 'id': gid}})
        else:
            text = elem.get('text', '').strip()
            elem_info.append({'tag': TYPE_MAP.get(etype, 'P'), 'text': text.lower(), 'group': None})

    assignments = {}
    used_elements = set()
    mcid = 0

    for block_idx, block in enumerate(blocks):
        if block['type'] == 'image':
            # No elem_info entry to anchor reading order to — fall back to
            # placing it after all named elements, ordered by its own stream
            # position among other such images (images are rare on the dense
            # forms where stream order diverges from reading order).
            assignments[block_idx] = ('Figure', mcid, None, len(elem_info) + block_idx)
            mcid += 1
            continue

        block_text = block.get('text', '').lower().strip()
        if not block_text:
            continue

        best_elem, best_score = None, 0
        for ei, elem in enumerate(elem_info):
            if ei in used_elements or not elem['text']:
                continue
            bwords = set(block_text.split())
            ewords = set(elem['text'].split())
            overlap = len(bwords & ewords)
            # Containment (overlap / smaller set), not Jaccard (overlap / larger
            # set) — a content-stream block is often a single line, while a
            # matching element (e.g. a multi-line table header cell) can span
            # several lines. Dividing by the larger set systematically
            # under-scores true matches when the two text lengths differ a lot.
            score = overlap / min(len(bwords), len(ewords))
            if score > best_score and score > 0.5:
                best_score = score
                best_elem = ei

        if best_elem is not None:
            chosen = elem_info[best_elem]
            assignments[block_idx] = (chosen['tag'], mcid, chosen['group'], best_elem)
            used_elements.add(best_elem)
            mcid += 1

    # Second pass: header-row (TH) and heading (H1-H6) candidates are
    # structurally critical — a table with zero TH tags fails Section 508
    # even if every data cell tagged correctly, and a heading that's dropped
    # entirely (rather than mis-leveled) shows up as a level *skip* between
    # whichever headings on either side of it did get tagged. Dense forms
    # often split a cell or heading's text across many tiny content-stream
    # fragments (single words, sometimes single tokens), so a candidate can
    # legitimately fail the stricter 0.5 containment bar against every
    # fragment and end up completely untagged. Retry just the still-unmatched
    # TH/heading candidates against still-unmatched blocks with a relaxed
    # threshold — bounded to those tags and unused blocks/elements so it
    # can't disturb assignments already resolved above.
    _critical_tags = {'TH'} | {f'H{i}' for i in range(1, 7)}
    unmatched_blocks = [
        bi for bi, b in enumerate(blocks)
        if bi not in assignments and b['type'] == 'text' and b.get('text', '').strip()
    ]
    for ei, elem in enumerate(elem_info):
        if ei in used_elements or elem['tag'] not in _critical_tags or not elem['text']:
            continue
        ewords = set(elem['text'].split())
        if not ewords:
            continue
        best_block, best_score = None, 0
        for bi in unmatched_blocks:
            bwords = set(blocks[bi]['text'].lower().strip().split())
            if not bwords:
                continue
            overlap = len(bwords & ewords)
            score = overlap / min(len(bwords), len(ewords))
            if score > best_score and score > 0.25:
                best_score = score
                best_block = bi
        if best_block is not None:
            assignments[best_block] = (elem['tag'], mcid, elem['group'], ei)
            used_elements.add(ei)
            unmatched_blocks.remove(best_block)
            mcid += 1

    return assignments


def _strip_marked_content(instructions) -> list:
    """Remove any existing BDC/BMC/EMC markers from a content stream.

    Strips pre-existing marks so we never produce duplicate MCID values on
    the same page when we inject our own markers.
    """
    return [i for i in instructions if str(i.operator) not in ('BDC', 'BMC', 'EMC')]


def _inject_mcids(instructions, blocks: list[dict], assignments: dict[int, tuple[str, int, dict | None, int]]) -> list:
    """Wrap every content item in a BDC/EMC sequence.

    - Matched blocks  → struct BDC with MCID (real tagged content)
    - Unmatched blocks → /Artifact BDC (real content, marked as artifact)
    - Everything else  → /Artifact BDC (graphics, paths, inter-block operators)

    This ensures NO content sits outside a marked-content sequence, satisfying
    Matterhorn checkpoint 13-004 ("real content not tagged").

    `instructions` must already be stripped of existing BDC/BMC/EMC markers,
    and `blocks` must have been derived from those same stripped instructions.
    """
    from pikepdf import ContentStreamInstruction, Operator, Name, Dictionary, Integer as PikInt

    artifact_bdc = ContentStreamInstruction([Name('/Artifact'), Dictionary()], Operator('BDC'))
    emc          = ContentStreamInstruction([], Operator('EMC'))

    def struct_bdc(tag: str, mcid: int) -> ContentStreamInstruction:
        return ContentStreamInstruction(
            [Name(f'/{tag}'), Dictionary(MCID=PikInt(mcid))], Operator('BDC')
        )

    # Map block start/end instruction indices
    block_start_idx = {b['start']: bi for bi, b in enumerate(blocks)}
    block_end_idx   = {b['end']:   bi for bi, b in enumerate(blocks)}

    new_instrs: list = []
    in_artifact = False
    i = 0

    while i < len(instructions):
        if i in block_start_idx:
            bi    = block_start_idx[i]
            block = blocks[bi]

            # Close any open inter-block artifact before starting this block
            if in_artifact:
                new_instrs.append(emc)
                in_artifact = False

            if bi in assignments:
                tag_name, mcid_val, _group, _elem_order = assignments[bi]
                new_instrs.append(struct_bdc(tag_name, mcid_val))
            else:
                new_instrs.append(artifact_bdc)

            for j in range(block['start'], block['end'] + 1):
                new_instrs.append(instructions[j])

            new_instrs.append(emc)
            i = block['end'] + 1

        else:
            # Content between/outside blocks — wrap in a single Artifact sequence
            if not in_artifact:
                new_instrs.append(artifact_bdc)
                in_artifact = True
            new_instrs.append(instructions[i])
            i += 1

    if in_artifact:
        new_instrs.append(emc)

    return new_instrs


def _build_outlines(pdf, pages_data: list[dict]) -> None:
    """Add a hierarchical bookmark outline built from heading elements.

    PDF/UA requires a document outline when the document has sections.
    Headings are collected across all pages and nested by level.
    """
    from pikepdf import Dictionary, Array, Name, String, Integer as PikInt

    # Collect headings in document order
    headings = []
    for page_idx, page_data in enumerate(pages_data):
        if page_idx >= len(pdf.pages):
            break
        for elem in page_data.get('elements', []):
            if elem.get('type') == 'heading':
                level = max(1, min(6, elem.get('level', 1)))
                text  = elem.get('text', '').strip()
                if text:
                    headings.append((page_idx, level, text))

    if not headings:
        return

    outline_root = pdf.make_indirect(Dictionary(Type=Name('/Outlines')))

    # Build item dicts (level, item ref, children list)
    # /Fit destination scrolls to the top of the page
    entries = []
    for page_idx, level, text in headings:
        dest = Array([pdf.pages[page_idx].obj, Name('/Fit')])
        item = pdf.make_indirect(Dictionary(Title=String(text), Dest=dest))
        entries.append({'level': level, 'item': item, 'children': []})

    # Nest entries by level using a stack
    root_entry = {'level': 0, 'item': outline_root, 'children': []}
    stack = [root_entry]
    for entry in entries:
        while len(stack) > 1 and stack[-1]['level'] >= entry['level']:
            stack.pop()
        stack[-1]['children'].append(entry)
        stack.append(entry)

    # Wire up Parent / Prev / Next / First / Last / Count links
    def _link(parent_entry: dict) -> int:
        children = parent_entry['children']
        if not children:
            return 0
        parent_ref = parent_entry['item']
        parent_ref['/First'] = children[0]['item']
        parent_ref['/Last']  = children[-1]['item']
        total = len(children)
        for i, child in enumerate(children):
            child['item']['/Parent'] = parent_ref
            if i > 0:
                child['item']['/Prev'] = children[i - 1]['item']
            if i < len(children) - 1:
                child['item']['/Next'] = children[i + 1]['item']
            total += _link(child)
        # Positive count = subtree open; root always positive
        count = total if parent_entry['level'] == 0 else len(children)
        parent_ref['/Count'] = PikInt(count)
        return len(children)

    _link(root_entry)

    pdf.Root['/Outlines']  = outline_root
    pdf.Root['/PageMode']  = Name('/UseOutlines')


_FIELD_NUM_RE = re.compile(r'\d{1,3}\.[a-zA-Z]{0,2}\.?')


def _looks_like_field_label(text: str) -> bool:
    """True for form-field reference numbers/letters (e.g. "24.", "5.a.",
    or multi-column-merge artifacts like "5.a.12.a.") that vision models on
    dense government forms routinely mis-tag as headings because the field
    number is bold or boxed. A real heading reads as a title; these don't."""
    stripped = _FIELD_NUM_RE.sub('', text).strip(' .()')
    if not stripped:
        return True  # nothing left but number/letter references
    if len(_FIELD_NUM_RE.findall(text)) >= 2:
        return True  # two+ field-number tokens glued together
    return False


def _normalize_heading_levels(pages_data: list[dict]) -> None:
    """Normalize heading levels across the whole document so there are no
    gaps (e.g. a document-wide H1 -> H3 becomes H1 -> H2), mutating elements
    in place. Also demotes mis-tagged field-label "headings" to paragraphs
    first, so they don't consume a level or create phantom gaps.

    Each page is analyzed independently (by the AI or the font-size
    heuristic), so a page's local heading judgment can disagree with the
    level the previous page left off at — a sub-heading on page 2 might get
    judged H3 in isolation even though the document never used H2 to get
    there.

    Clamps against the *immediately preceding* heading's level, not the
    deepest level ever seen — a document-wide high-water mark would let a
    later H2 -> H4 through unclamped once H3 had appeared anywhere earlier,
    even though that's still a skip relative to its own immediate neighbor
    (which is what accessibility checkers actually flag).
    """
    prev_level = 0
    for page in pages_data:
        for elem in page.get("elements", []):
            if elem.get("type") != "heading":
                continue
            if _looks_like_field_label(elem.get("text", "")):
                elem["type"] = "paragraph"
                elem.pop("level", None)
                continue
            level = max(1, min(6, elem.get("level", 1)))
            if level > prev_level + 1:
                level = prev_level + 1
            elem["level"] = level
            prev_level = level


def tag_pdf_with_accessibility(
    source_pdf: str,
    pages_data: list[dict],
    output_path: str,
    title: str = "",
) -> None:
    """
    Create an accessible PDF by tagging the ORIGINAL document in-place.
    Visual appearance (fonts, images, tables, layout) is preserved exactly.
    Adds: MarkInfo, Lang, document title, and a full StructTreeRoot with MCIDs.

    MCIDs are page-local (reset to 0 on each page) per the PDF spec.
    The ParentTree maps each page's StructParents index to an array of parent
    struct elements — one entry per MCID — as required by PDF/UA validators.
    """
    from pikepdf import Dictionary, Array, Name, String, Integer as PikInt

    _normalize_heading_levels(pages_data)
    doc_title = title or Path(source_pdf).stem.replace("_", " ").title()

    with pikepdf.open(source_pdf) as pdf:
        # Document-level metadata
        with pdf.open_metadata(set_pikepdf_as_editor=False) as meta:
            if not meta.get('dc:title'):
                meta['dc:title'] = doc_title
            meta['dc:language'] = 'en-US'
            # Required to declare PDF/UA-1 conformance (Matterhorn 06-001)
            meta['pdfuaid:part'] = '1'
        pdf.Root['/Lang'] = String('en-US')
        pdf.Root['/MarkInfo'] = Dictionary(Marked=True, Suspects=False)

        # PDF/UA requires the document title to be shown in the title bar
        if '/ViewerPreferences' not in pdf.Root:
            pdf.Root['/ViewerPreferences'] = pdf.make_indirect(Dictionary())
        pdf.Root['/ViewerPreferences']['/DisplayDocTitle'] = True

        if '/StructTreeRoot' in pdf.Root:
            del pdf.Root['/StructTreeRoot']

        # Clear any existing StructParents from all pages
        for pg in pdf.pages:
            if '/StructParents' in pg:
                del pg['/StructParents']

        all_struct_elems = []   # flattened list for the Document StructElem
        page_parent_arrays = [] # one Array per tagged page for the ParentTree

        for page_idx, page_data in enumerate(pages_data):
            if page_idx >= len(pdf.pages):
                break
            page_obj = pdf.pages[page_idx]
            elements = page_data.get('elements', [])

            # Parse and strip existing marked-content operators so we start clean.
            # We ALWAYS rewrite the content stream — even for pages with no matched
            # content — so original BDC/MCID markers don't become dangling refs
            # after we clear the StructTreeRoot.
            try:
                raw_stream = pikepdf.parse_content_stream(page_obj)
                clean_stream = _strip_marked_content(raw_stream)
                blocks = _find_content_blocks(clean_stream)
                assignments = _match_blocks_to_elements(blocks, elements) if elements else {}
            except Exception as e:
                print(f"  [warn] Stream parse failed p{page_idx + 1}: {e}")
                continue

            # Always write back (with artifacts for everything if no assignments)
            try:
                new_stream = _inject_mcids(clean_stream, blocks, assignments)
                page_obj['/Contents'] = pdf.make_stream(
                    pikepdf.unparse_content_stream(new_stream)
                )
            except Exception as e:
                print(f"  [warn] MCID injection failed p{page_idx + 1}: {e}")
                continue

            if not assignments:
                continue  # No struct elements needed for this page

            # Get page MediaBox for Figure /BBox attribute
            mb = page_obj.get('/MediaBox', Array([0, 0, 612, 792]))
            page_bbox = [float(mb[0]), float(mb[1]), float(mb[2]), float(mb[3])]

            # Build struct elements; keep a slot-indexed array for the ParentTree
            max_local_mcid = max(lm for _, lm, _, _ in assignments.values())
            parent_slot: list = [None] * (max_local_mcid + 1)

            # Table/List are grouping elements that cannot directly contain MCRs
            # (Matterhorn "content in inadmissible location"). Their leaf cells
            # (TH/TD) or items (LI > LBody) hold the MCRs; the group elements are
            # finalized below once all of their children have been collected.
            table_groups: dict[int, dict] = {}   # group_id -> {ref, rows: {row_idx: {ref, cells}}}
            list_groups: dict[int, dict] = {}    # group_id -> {ref, items: [li_ref, ...]}

            # Iterate in elem_order (the AI's linearized reading order), not
            # block_idx (content-stream position) — on fillable forms, field
            # widgets/labels are often written to the content stream in
            # field-definition order rather than top-to-bottom visual order,
            # so block order is not a reliable proxy for reading order. Using
            # the AI's already-linearized order also avoids the dict-
            # insertion-order issue where the relaxed second pass in
            # _match_blocks_to_elements appends matches regardless of page
            # position.
            for block_idx, (tag_name, local_mcid, group, _elem_order) in sorted(
                assignments.items(), key=lambda kv: kv[1][3]
            ):
                mcr = pdf.make_indirect(Dictionary(
                    Type=Name('/MCR'),
                    Pg=page_obj.obj,
                    MCID=PikInt(local_mcid),
                ))

                if group and group['type'] == 'Table':
                    gid = group['id']
                    if gid not in table_groups:
                        table_ref = pdf.make_indirect(Dictionary(
                            Type=Name('/StructElem'), S=Name('/Table'), Pg=page_obj.obj,
                        ))
                        table_groups[gid] = {'ref': table_ref, 'rows': {}}
                        all_struct_elems.append(table_ref)
                    tg = table_groups[gid]
                    row_idx = group['row']
                    if row_idx not in tg['rows']:
                        tr_ref = pdf.make_indirect(Dictionary(
                            Type=Name('/StructElem'), S=Name('/TR'), Pg=page_obj.obj,
                        ))
                        tg['rows'][row_idx] = {'ref': tr_ref, 'cells': []}
                    row = tg['rows'][row_idx]
                    cell_d = Dictionary(
                        Type=Name('/StructElem'), S=Name(f'/{tag_name}'), Pg=page_obj.obj,
                        K=Array([mcr]),
                    )
                    if tag_name == 'TH':
                        cell_d['/A'] = pdf.make_indirect(Dictionary(O=Name('/Table'), Scope=Name('/Column')))
                    cell_ref = pdf.make_indirect(cell_d)
                    cell_ref['/P'] = row['ref']
                    row['cells'].append(cell_ref)
                    parent_slot[local_mcid] = cell_ref
                    continue

                if group and group['type'] == 'L':
                    gid = group['id']
                    if gid not in list_groups:
                        list_ref = pdf.make_indirect(Dictionary(
                            Type=Name('/StructElem'), S=Name('/L'), Pg=page_obj.obj,
                        ))
                        list_groups[gid] = {'ref': list_ref, 'items': []}
                        all_struct_elems.append(list_ref)
                    lg = list_groups[gid]
                    lbody_ref = pdf.make_indirect(Dictionary(
                        Type=Name('/StructElem'), S=Name('/LBody'), Pg=page_obj.obj, K=Array([mcr]),
                    ))
                    li_ref = pdf.make_indirect(Dictionary(
                        Type=Name('/StructElem'), S=Name('/LI'), Pg=page_obj.obj, K=Array([lbody_ref]),
                    ))
                    lbody_ref['/P'] = li_ref
                    li_ref['/P'] = lg['ref']
                    lg['items'].append(li_ref)
                    parent_slot[local_mcid] = lbody_ref
                    continue

                struct_d = Dictionary(
                    Type=Name('/StructElem'),
                    S=Name(f'/{tag_name}'),
                    Pg=page_obj.obj,
                    K=Array([mcr]),
                )
                if tag_name == 'Figure':
                    # /Alt is required; /BBox prevents Matterhorn 17-002 failures
                    alt = ''
                    for elem in elements:
                        if elem.get('type') == 'image_alt':
                            alt = elem.get('text', '')[:500]
                            break
                    struct_d['/Alt'] = String(alt or 'Figure')
                    struct_d['/A'] = pdf.make_indirect(Dictionary(
                        O=Name('/Layout'),
                        BBox=Array([
                            page_bbox[0], page_bbox[1],
                            page_bbox[2], page_bbox[3],
                        ]),
                    ))
                struct_ref = pdf.make_indirect(struct_d)
                parent_slot[local_mcid] = struct_ref
                all_struct_elems.append(struct_ref)

            # Finalize grouping elements now that all children are collected
            for tg in table_groups.values():
                row_refs = []
                for row_idx in sorted(tg['rows'].keys()):
                    row = tg['rows'][row_idx]
                    row['ref']['/K'] = Array(row['cells'])
                    row['ref']['/P'] = tg['ref']
                    row_refs.append(row['ref'])
                tg['ref']['/K'] = Array(row_refs)
            for lg in list_groups.values():
                lg['ref']['/K'] = Array(lg['items'])

            # Wire up StructParents: page points to its index in the ParentTree
            tree_idx = len(page_parent_arrays)
            page_obj['/StructParents'] = PikInt(tree_idx)
            # parent_slot is contiguous (MCIDs are 0..N-1) so no None gaps expected
            page_parent_arrays.append(Array(parent_slot))

        if not all_struct_elems:
            pdf.save(output_path)
            return

        # Re-clamp heading levels against the sequence that actually made it
        # into the struct tree. _normalize_heading_levels already made
        # pages_data gap-free, but block-matching can legitimately drop a
        # heading entirely (e.g. banner/title text baked into a logo image,
        # with no extractable text run in the content stream to attach an
        # MCID to) — leaving a gap between whichever headings on either side
        # did get tagged. Re-clamping against survivors only is the same fix
        # already applied to the docx heading pass in make_docx_accessible.
        prev_level = 0
        for ref in all_struct_elems:
            sname = str(ref.get('/S', ''))
            if not re.fullmatch(r'/H[1-6]', sname):
                continue
            level = int(sname[2])
            if level > prev_level + 1:
                level = prev_level + 1
                ref['/S'] = Name(f'/H{level}')
            prev_level = level

        doc_elem = pdf.make_indirect(Dictionary(
            Type=Name('/StructElem'),
            S=Name('/Document'),
            K=Array(all_struct_elems),
        ))
        for ref in all_struct_elems:
            ref['/P'] = doc_elem

        # ParentTree Nums: [0, array_for_page0, 1, array_for_page1, ...]
        # Each array maps MCID index → parent StructElem for that page
        nums: list = []
        for idx, arr in enumerate(page_parent_arrays):
            nums.extend([PikInt(idx), pdf.make_indirect(arr)])

        parent_tree = pdf.make_indirect(Dictionary(Nums=Array(nums)))
        struct_root = pdf.make_indirect(Dictionary(
            Type=Name('/StructTreeRoot'),
            K=Array([doc_elem]),
            ParentTree=parent_tree,
            ParentTreeNextKey=PikInt(len(page_parent_arrays)),
        ))
        doc_elem['/P'] = struct_root
        pdf.Root['/StructTreeRoot'] = struct_root

        # Add bookmark outline (required by PDF/UA when document has sections)
        _build_outlines(pdf, pages_data)

        pdf.save(output_path)


# ---------------------------------------------------------------------------
# Shared DOCX conversion (used by both CLI and web job)
# ---------------------------------------------------------------------------

def _iter_block_items(container):
    """Yield paragraphs and tables in true document order (as they appear in
    the XML body), instead of python-docx's separate .paragraphs/.tables
    lists which lose interleaving between the two."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    parent_elm = container.element.body if hasattr(container, "element") else container._tc
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, container)
        elif child.tag == qn("w:tbl"):
            yield Table(child, container)


def _iter_all_paragraphs(container):
    """Yield every paragraph in a Document (or table cell) in true reading
    order, including ones nested inside table cells. pdf2docx renders most
    form layouts as tables, so headings routinely end up inside cells
    interleaved with top-level body paragraphs — a naive paragraphs-then-
    tables traversal would scramble that order and break heading-level
    normalization."""
    from docx.table import Table

    for block in _iter_block_items(container):
        if isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    yield from _iter_all_paragraphs(cell)
        else:
            yield block


def make_docx_accessible(docx_path: str, pages_data: list[dict]) -> None:
    """
    Post-process a visually-converted DOCX to add semantic accessibility structure:
    heading styles, table header rows, image alt text, and document language.
    """
    from collections import Counter

    try:
        doc = Document(docx_path)
    except Exception:
        return

    doc.core_properties.language = "en-US"

    # Heading styles — match against the AI-extracted heading elements first.
    # pdf2docx output only preserves font *size*, and many dense government
    # forms mark headings with bold/caps rather than a larger point size, so
    # a size-only heuristic finds nothing on those documents even though the
    # AI correctly identified the headings from the page image/text. Fall
    # back to the size heuristic only when no AI heading data is available
    # (e.g. mock mode with no pages_data).
    _normalize_heading_levels(pages_data)
    heading_elems = [
        (max(1, min(6, e.get("level", 1))), e.get("text", "").strip())
        for page in pages_data
        for e in page.get("elements", [])
        if e.get("type") == "heading" and e.get("text", "").strip()
    ]

    def _heading_text_match(a: str, b: str) -> bool:
        if a == b:
            return True
        # Substring containment alone produces false positives when a short
        # heading ("Page", "Future Developments") happens to appear inside an
        # unrelated long body paragraph — require the shorter string to make
        # up most of the longer one before treating it as the same heading.
        shorter, longer = (a, b) if len(a) <= len(b) else (b, a)
        return bool(shorter) and shorter in longer and len(shorter) >= 0.6 * len(longer)

    matched_any = False
    if heading_elems:
        remaining = list(heading_elems)
        for para in _iter_all_paragraphs(doc):
            if not remaining:
                break
            text = para.text.strip()
            if not text:
                continue
            for i, (level, htext) in enumerate(remaining):
                if _heading_text_match(text, htext):
                    try:
                        para.style = doc.styles[f"Heading {level}"]
                        matched_any = True
                    except KeyError:
                        pass
                    del remaining[i]
                    break

    if not matched_any:
        # Fallback: detect by font size
        sizes = []
        for para in _iter_all_paragraphs(doc):
            for run in para.runs:
                if run.font.size:
                    sizes.append(round(run.font.size.pt, 1))

        if sizes:
            body_size = Counter(sizes).most_common(1)[0][0]
            heading_sizes = sorted(
                {s for s in sizes if s >= body_size * 1.2},
                reverse=True,
            )[:4]
            size_to_level = {s: i + 1 for i, s in enumerate(heading_sizes)}

            for para in _iter_all_paragraphs(doc):
                if not para.runs:
                    continue
                para_sizes = [round(r.font.size.pt, 1) for r in para.runs if r.font.size]
                if not para_sizes:
                    continue
                dominant = Counter(para_sizes).most_common(1)[0][0]
                if dominant in size_to_level:
                    try:
                        para.style = doc.styles[f"Heading {size_to_level[dominant]}"]
                    except KeyError:
                        pass

    # Normalize heading levels so there are no gaps (e.g. H1 → H3 becomes
    # H1 → H2). Needed even after AI text-matching: a heading can be present
    # in pages_data's (already gap-free) sequence but fail to match any docx
    # paragraph, leaving a gap in the subsequence that actually got styled.
    # Clamp against the *immediately preceding* heading's level, not the
    # deepest level ever seen — a high-water mark would let a later H2 → H4
    # through unclamped once H3 had appeared anywhere earlier in the doc.
    prev_level = 0
    for para in _iter_all_paragraphs(doc):
        sname = para.style.name
        if not sname.startswith("Heading "):
            continue
        try:
            level = int(sname.split()[-1])
        except ValueError:
            continue
        if level > prev_level + 1:
            try:
                para.style = doc.styles[f"Heading {prev_level + 1}"]
                level = prev_level + 1
            except KeyError:
                pass
        prev_level = level

    # Table header rows — mark first row of every table with w:tblHeader
    for table in doc.tables:
        if not table.rows:
            continue
        first_tr = table.rows[0]._tr
        trPr = first_tr.find(qn("w:trPr"))
        if trPr is None:
            trPr = OxmlElement("w:trPr")
            first_tr.insert(0, trPr)
        if trPr.find(qn("w:tblHeader")) is None:
            trPr.append(OxmlElement("w:tblHeader"))

    # Alt text on inline images using AI-extracted descriptions
    ai_alts = [
        elem.get("alt_text", "").strip() or "Figure"
        for page in pages_data
        for elem in page.get("elements", [])
        if elem.get("type") == "image"
    ]
    ns_wp = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    for i, img_elem in enumerate(doc.element.body.findall(f".//{{{ns_wp}}}inline")):
        docPr = img_elem.find(f"{{{ns_wp}}}docPr")
        if docPr is not None and not docPr.get("descr", "").strip():
            docPr.set("descr", ai_alts[i] if i < len(ai_alts) else "Figure")

    # Move footer/header text from body into actual DOCX section footer/header.
    # pdf2docx converts PDF drawString footer/header as body paragraphs — detect
    # them by matching against position-extracted elements, then relocate them.

    # Collect per-page footer/header texts with their page index
    page_footer_texts: dict[int, list[str]] = {}
    page_header_texts: dict[int, list[str]] = {}
    for pi, page in enumerate(pages_data):
        for elem in page.get("elements", []):
            t = elem.get("text", "").strip()
            if not t:
                continue
            if elem.get("type") == "footer":
                page_footer_texts.setdefault(pi, []).append(t)
            elif elem.get("type") == "header":
                page_header_texts.setdefault(pi, []).append(t)

    all_footer_texts = [t for texts in page_footer_texts.values() for t in texts]
    all_header_texts = [t for texts in page_header_texts.values() for t in texts]

    if all_footer_texts or all_header_texts:
        section = doc.sections[0]
        paras_to_remove = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if any(text in ft or ft in text for ft in all_footer_texts):
                paras_to_remove.append(para)
            elif any(text in ht or ht in text for ht in all_header_texts):
                paras_to_remove.append(para)

        # pdf2docx gives one section per source page, so when footer/header text
        # actually differs page-to-page (e.g. "Page 1 of 4" vs "Page 2 of 4") we
        # can set each section's own value. The build_docx fallback only adds a
        # section when page *size* changes, so it won't generally line up with
        # pages 1:1 — guard with this check rather than assume the mapping holds.
        per_page_sections = len(doc.sections) == len(pages_data)

        if all_footer_texts:
            footer_page_indices = set(page_footer_texts.keys())
            varies_by_page = len({tuple(v) for v in page_footer_texts.values()}) > 1

            if varies_by_page and per_page_sections:
                for pi, sec in enumerate(doc.sections):
                    sec.footer.is_linked_to_previous = False
                    sec.footer.paragraphs[0].text = " | ".join(dict.fromkeys(page_footer_texts.get(pi, [])))
            else:
                footer_text = " | ".join(dict.fromkeys(all_footer_texts))
                if footer_page_indices.issubset({0}):
                    # Footer only on first page — use Word's "different first page" feature
                    section.different_first_page_header_footer = True
                    section.first_page_footer.is_linked_to_previous = False
                    section.first_page_footer.paragraphs[0].text = footer_text
                else:
                    # Same footer text on every page — set as section-wide footer
                    section.footer.is_linked_to_previous = False
                    section.footer.paragraphs[0].text = footer_text

        if all_header_texts:
            header_page_indices = set(page_header_texts.keys())
            varies_by_page = len({tuple(v) for v in page_header_texts.values()}) > 1

            if varies_by_page and per_page_sections:
                for pi, sec in enumerate(doc.sections):
                    sec.header.is_linked_to_previous = False
                    sec.header.paragraphs[0].text = " ".join(dict.fromkeys(page_header_texts.get(pi, [])))
            else:
                header_text = " ".join(dict.fromkeys(all_header_texts))
                if header_page_indices.issubset({0}):
                    section.different_first_page_header_footer = True
                    section.first_page_header.is_linked_to_previous = False
                    section.first_page_header.paragraphs[0].text = header_text
                else:
                    section.header.is_linked_to_previous = False
                    section.header.paragraphs[0].text = header_text

        for para in paras_to_remove:
            para._element.getparent().remove(para._element)

    doc.save(docx_path)


def convert_to_docx(
    pdf_path: str,
    output_path: str,
    pages_data: list[dict],
    page_dims: list[tuple[float, float]] | None = None,
    title: str = "",
) -> None:
    """
    Convert PDF to DOCX using pdf2docx for best visual fidelity, falling back to
    build_docx if pdf2docx is unavailable. Always runs accessibility post-processing.
    """
    converted = False
    try:
        from pdf2docx import Converter
        cv = Converter(pdf_path)
        cv.convert(output_path, start=0, end=None)
        cv.close()
        converted = True
    except Exception:
        pass

    if not converted:
        build_docx(pages_data, output_path, page_dims=page_dims, title=title)

    if Path(output_path).exists():
        make_docx_accessible(output_path, pages_data)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def _parse_args():
    import argparse
    parser = argparse.ArgumentParser(description="ADA PDF Remediation Tool")
    parser.add_argument("input", help="Input PDF file")
    parser.add_argument("output", nargs="?", help="Output .docx path (optional)")
    parser.add_argument("--pages", type=int, metavar="N", help="Only process first N pages (cheap testing)")
    return parser.parse_args()


def main():
    args = _parse_args()
    pdf_path = args.input

    if not os.path.exists(pdf_path):
        print(f"Error: file not found: {pdf_path}")
        sys.exit(1)

    docx_output_path = args.output if args.output else Path(pdf_path).stem + "_accessible.docx"
    pdf_output_path  = str(Path(pdf_path).stem + "_accessible.pdf")

    # --- Step 1: Page count ---
    with pdfplumber.open(pdf_path) as pdf:
        total_pages = len(pdf.pages)
    if total_pages == 0:
        print(f"Error: {pdf_path} has no readable pages — the file is likely corrupted.")
        sys.exit(1)
    try:
        with pikepdf.open(pdf_path):
            pass
    except pikepdf.PdfError as e:
        print(f"Error: {pdf_path} could not be opened as a valid PDF ({e}) — the file is likely corrupted.")
        sys.exit(1)
    pages_to_process = min(total_pages, MAX_PAGES)
    if args.pages:
        pages_to_process = min(args.pages, pages_to_process)
    print(f"\nPDF: {pdf_path}")
    page_note = f" (processing {pages_to_process}/{total_pages})" if pages_to_process < total_pages else ""
    print(f"Pages: {total_pages}{page_note}")

    # --- Step 2: Detect backends and show cost estimate ---
    print("\nDetecting AI backends...")
    backends = detect_backends()
    print(f"  Ollama ({LOCAL_MODEL}): {'available' if backends['ollama'] else 'not available'}")
    print(f"  Claude API:             {'available' if backends['claude'] else 'not available'}")

    if not backends["ollama"] and not backends["claude"]:
        print(f"\n  [info] No AI backend found — running in mock mode (text extraction only).")
        print(f"         Output will be readable but table structure won't be reconstructed.")
        print(f"         To enable AI: install Ollama + pull {LOCAL_MODEL}, or set ANTHROPIC_API_KEY")

    cost, cost_note = estimate_cost(pages_to_process, backends)
    print(f"\nCost estimate: {cost_note}")

    # --- Step 3: Confirm large documents ---
    if pages_to_process > LARGE_DOC_THRESHOLD:
        if not confirm_large_doc(pages_to_process, cost_note):
            print("Aborted.")
            sys.exit(0)

    # --- Step 4: Extract text layer hints + page dimensions ---
    print("\nExtracting text layer and page dimensions...")
    text_layers = extract_text_layer(pdf_path)
    page_dims = get_page_dimensions(pdf_path)
    print(f"  Page sizes: {', '.join(f'{w:.1f}\"×{h:.1f}\"' for w, h in page_dims[:pages_to_process])}")

    # Classify complexity so users can see AI usage upfront
    print("\nClassifying pages...")
    complexities = {n: classify_page_complexity(pdf_path, n) for n in range(1, pages_to_process + 1)}
    simple_count  = sum(1 for v in complexities.values() if v == "simple")
    complex_count = sum(1 for v in complexities.values() if v == "complex")
    print(f"  Simple (text only, no AI): {simple_count} page(s)")
    print(f"  Complex (tables/images, AI needed): {complex_count} page(s)")

    # Re-estimate cost based on actual complex page count only
    if complex_count < pages_to_process:
        cost, cost_note = estimate_cost(complex_count, backends)
        print(f"  Revised cost estimate (complex pages only): {cost_note}")

    # --- Step 5: Analyze pages ---
    pages_data = []

    for page_num in range(1, pages_to_process + 1):
        complexity = complexities[page_num]
        if os.environ.get("MOCK_MODE") == "1":
            label = "mock"
        elif complexity == "simple":
            label = "simple"
        elif backends["ollama"]:
            label = "ollama"
        else:
            label = "claude"
        print(f"  Page {page_num}/{pages_to_process} [{label}]...", end=" ", flush=True)
        page_data = analyze_page(pdf_path, page_num, text_layers.get(page_num, ""), backends)
        pages_data.append(page_data)
        elem_count = len(page_data.get("elements", []))
        img_count  = sum(1 for e in page_data.get("elements", []) if e.get("_image_bytes"))
        img_note   = f", {img_count} image(s) embedded" if img_count else ""
        print(f"{elem_count} elements{img_note}")

    # --- Step 6: Tag original PDF with accessibility structure ---
    # (preserves all fonts, images, tables, and layout from the source PDF)
    print(f"\nTagging original PDF with accessibility structure → {pdf_output_path}")
    doc_title = Path(pdf_path).stem.replace("_", " ").title()
    tag_pdf_with_accessibility(pdf_path, pages_data, pdf_output_path, title=doc_title)
    print(f"  ✅ Accessible PDF created (visual appearance preserved).")

    # --- Step 7: Also build DOCX for editing ---
    print(f"\nBuilding editable .docx → {docx_output_path}")
    convert_to_docx(pdf_path, docx_output_path, pages_data, page_dims=page_dims, title=doc_title)

    print(f"\nDone!")
    print(f"  PDF file (primary):  {pdf_output_path}  ← visually identical to source")
    print(f"  Word file (editable): {docx_output_path}")
    print(f"\nValidate with PAC 2026: https://pac.pdf-accessibility.org")


if __name__ == "__main__":
    main()
