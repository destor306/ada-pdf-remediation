"""
Job queue — auto-selects Redis/RQ when REDIS_URL is set, falls back to in-memory threads.

When REDIS_URL is configured:
  - Job state is stored in Redis so web-process and worker-process stay in sync.
  - Jobs are executed by RQ workers (run: rq worker ada).
When REDIS_URL is absent:
  - Jobs run in daemon threads within the web process (dev/single-server mode).
"""

import uuid
import json
import threading
import traceback
from datetime import datetime, timezone
from pathlib import Path
from typing import Literal

from app.config import UPLOAD_DIR, OUTPUT_DIR, REDIS_URL

JobStatus = Literal["queued", "running", "done", "failed"]

_REDIS_TTL = 86400 * 7  # keep job state 7 days


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


class Job:
    def __init__(
        self,
        job_id: str,
        pdf_path: str,
        output_path: str,
        use_claude: bool = False,
        notify_email: str = "",
        user_id: int | None = None,
        original_stem: str = "",
    ):
        self.id            = job_id
        self.pdf_path      = pdf_path
        self.output_path   = output_path
        self.output_pdf    = output_path.replace(".docx", "_ada.pdf")
        self.use_claude    = use_claude
        self.notify_email  = notify_email
        self.user_id       = user_id
        self.original_stem = original_stem or Path(pdf_path).stem
        self.status: JobStatus = "queued"
        self.progress      = 0
        self.current_page  = 0
        self.total_pages   = 0
        self.error: str | None = None
        self.check_report: dict | None = None
        self.backend       = "unknown"
        self.created_at    = _now()
        self.completed_at: str | None = None

    def to_dict(self) -> dict:
        return {k: v for k, v in self.__dict__.items()}

    @classmethod
    def from_dict(cls, d: dict) -> "Job":
        obj = cls.__new__(cls)
        obj.__dict__.update(d)
        return obj


# ---------- in-memory fallback ----------
_jobs: dict[str, Job] = {}
_lock = threading.Lock()


# ---------- Redis helpers ----------

def _redis():
    if not REDIS_URL:
        return None
    try:
        import redis
        return redis.from_url(REDIS_URL)
    except Exception:
        return None


def _r_key(job_id: str) -> str:
    return f"accessifix:job:{job_id}"


def _r_index_key() -> str:
    return "accessifix:jobs"


def _store(job: Job):
    with _lock:
        _jobs[job.id] = job
    r = _redis()
    if r:
        r.setex(_r_key(job.id), _REDIS_TTL, json.dumps(job.to_dict()))
        r.sadd(_r_index_key(), job.id)
        r.expire(_r_index_key(), _REDIS_TTL)


def get_job(job_id: str) -> Job | None:
    r = _redis()
    if r:
        raw = r.get(_r_key(job_id))
        if raw:
            return Job.from_dict(json.loads(raw))
    with _lock:
        return _jobs.get(job_id)


def list_jobs() -> list[Job]:
    r = _redis()
    if r:
        ids = r.smembers(_r_index_key())
        jobs = []
        for jid in ids:
            raw = r.get(_r_key(jid.decode() if isinstance(jid, bytes) else jid))
            if raw:
                jobs.append(Job.from_dict(json.loads(raw)))
        return jobs
    with _lock:
        return list(_jobs.values())


# ---------- helpers ----------

def _convert_with_libreoffice(pdf_path: str, output_docx: str) -> bool:
    """
    Convert pdf_path → output_docx using LibreOffice headless.
    Uses a per-call user-profile temp dir so concurrent jobs don't collide.
    Returns True on success, False if LibreOffice is unavailable or fails.
    """
    import shutil
    import subprocess
    import tempfile

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return False

    tmp_profile = tempfile.mkdtemp(prefix="lo_profile_")
    try:
        out_dir = Path(output_docx).parent
        subprocess.run(
            [
                soffice,
                "--headless",
                f"-env:UserInstallation=file://{tmp_profile}",
                "--infilter=writer_pdf_import",
                "--convert-to", "docx",
                "--outdir", str(out_dir),
                pdf_path,
            ],
            capture_output=True,
            timeout=180,
        )
        # LibreOffice exits non-zero on Java warnings even when conversion succeeds,
        # so we check for the output file rather than trusting the return code.
        lo_out = out_dir / (Path(pdf_path).stem + ".docx")
        if not lo_out.exists():
            return False
        lo_out.rename(output_docx)
        return True
    except Exception:
        return False
    finally:
        shutil.rmtree(tmp_profile, ignore_errors=True)


def _make_docx_accessible(docx_path: str, pages_data: list) -> None:
    """
    Post-process a visually-converted DOCX to add semantic accessibility structure:
    - Document language declaration (en-US)
    - Word Heading styles derived from font-size heuristics
    - w:tblHeader on every table's first row
    - Alt text (descr) on inline images, using AI-extracted descriptions when available
    """
    from collections import Counter
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    try:
        doc = Document(docx_path)
    except Exception:
        return

    # 1. Language
    doc.core_properties.language = "en-US"

    # 2. Heading styles — detect by font size
    sizes = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.size:
                sizes.append(round(run.font.size.pt, 1))

    if sizes:
        body_size = Counter(sizes).most_common(1)[0][0]
        # Only treat sizes meaningfully larger than body as headings (≥20% bigger),
        # and cap at H4 so generated levels are always sequential.
        heading_sizes = sorted(
            {s for s in sizes if s >= body_size * 1.2},
            reverse=True,
        )[:4]
        size_to_level = {s: i + 1 for i, s in enumerate(heading_sizes)}

        for para in doc.paragraphs:
            if not para.runs:
                continue
            para_sizes = [round(r.font.size.pt, 1) for r in para.runs if r.font.size]
            if not para_sizes:
                continue
            dominant = Counter(para_sizes).most_common(1)[0][0]
            if dominant in size_to_level:
                level = size_to_level[dominant]
                style_name = f"Heading {level}"
                try:
                    para.style = doc.styles[style_name]
                except KeyError:
                    pass

        # Normalize heading levels: fill any gaps so H2→H4 becomes H2→H3
        max_seen = 0
        for para in doc.paragraphs:
            sname = para.style.name
            if not sname.startswith("Heading "):
                continue
            try:
                level = int(sname.split()[-1])
            except ValueError:
                continue
            if level > max_seen + 1:
                new_level = max_seen + 1
                try:
                    para.style = doc.styles[f"Heading {new_level}"]
                    level = new_level
                except KeyError:
                    pass
            max_seen = max(max_seen, level)

    # 3. Table header rows
    for table in doc.tables:
        if not table.rows:
            continue
        first_tr = table.rows[0]._tr
        trPr = first_tr.find(qn("w:trPr"))
        if trPr is None:
            trPr = OxmlElement("w:trPr")
            first_tr.insert(0, trPr)
        if trPr.find(qn("w:tblHeader")) is None:
            tblHeader = OxmlElement("w:tblHeader")
            trPr.append(tblHeader)

    # 4. Alt text on inline images
    ai_alts = []
    for page in pages_data:
        for elem in page.get("elements", []):
            if elem.get("type") == "image":
                ai_alts.append(elem.get("alt", "").strip() or "Figure")

    ns_wp = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    for i, img_elem in enumerate(doc.element.body.findall(f".//{{{ns_wp}}}inline")):
        docPr = img_elem.find(f"{{{ns_wp}}}docPr")
        if docPr is not None and not docPr.get("descr", "").strip():
            docPr.set("descr", ai_alts[i] if i < len(ai_alts) else "Figure")

    doc.save(docx_path)


# ---------- execution ----------

def _execute(job_id: str):
    """Worker entry point — loads job from store, runs pipeline, persists progress."""
    import sys
    sys.path.insert(0, str(Path(__file__).parent.parent))

    job = get_job(job_id)
    if job is None:
        return

    def _save():
        _store(job)

    job.status = "running"
    _save()
    try:
        import pdfplumber
        with pdfplumber.open(job.pdf_path) as pdf:
            job.total_pages = len(pdf.pages)
        _save()

        from ada_remediate import (
            detect_backends, extract_text_layer, get_page_dimensions,
            analyze_page, build_docx, tag_pdf_with_accessibility, MAX_PAGES,
        )

        backends = detect_backends()
        if not job.use_claude:
            backends["claude"] = False

        job.backend = "ollama" if backends["ollama"] else ("claude" if backends["claude"] else "mock")
        _save()

        pages_to_process = min(job.total_pages, MAX_PAGES)
        text_layers = extract_text_layer(job.pdf_path)
        page_dims   = get_page_dimensions(job.pdf_path)

        pages_data = []
        for page_num in range(1, pages_to_process + 1):
            job.current_page = page_num
            job.progress     = int((page_num - 1) / pages_to_process * 88)
            _save()
            pages_data.append(
                analyze_page(job.pdf_path, page_num, text_layers.get(page_num, ""), backends)
            )

        job.progress = 90
        _save()
        doc_title = Path(job.pdf_path).stem.replace("_", " ").title()
        tag_pdf_with_accessibility(job.pdf_path, pages_data, job.output_pdf, title=doc_title)

        job.progress = 93
        _save()
        try:
            from pdf2docx import Converter
            cv = Converter(job.pdf_path)
            cv.convert(job.output_path, start=0, end=None)
            cv.close()
        except Exception:
            build_docx(pages_data, job.output_path, page_dims=page_dims)

        if Path(job.output_path).exists():
            _make_docx_accessible(job.output_path, pages_data)

        job.progress = 97
        _save()
        try:
            from ada_check import CheckReport, run_docx_checks, run_verapdf, check_visual_similarity
            rpt = CheckReport(source_pdf=job.pdf_path, docx_path=job.output_path, pdf_path=job.output_pdf)
            run_docx_checks(job.pdf_path, job.output_path, rpt)
            if job.output_pdf:
                run_verapdf(job.output_pdf, rpt)
                check_visual_similarity(job.pdf_path, job.output_pdf, rpt)
            job.check_report = {
                "issues": [{"severity": i.severity, "category": i.category, "message": i.message} for i in rpt.issues],
                "passed": rpt.passed,
                "failed": rpt.failed,
            }
        except Exception as ce:
            job.check_report = {"error": str(ce)}

        job.progress     = 100
        job.status       = "done"
        job.completed_at = _now()
        _save()

        if job.user_id:
            try:
                from app.database import SessionLocal
                from app.models import User
                from app.auth import current_month
                db = SessionLocal()
                user = db.get(User, job.user_id)
                if user:
                    month = current_month()
                    if user.usage_month != month:
                        user.pages_used = 0
                        user.usage_month = month
                    user.pages_used = (user.pages_used or 0) + job.total_pages
                    db.commit()
            except Exception:
                pass
            finally:
                db.close()

        if job.notify_email:
            from app.email_notify import notify_done
            from app.config import APP_URL
            notify_done(job.notify_email, job.id, job.total_pages, APP_URL)

    except Exception:
        job.status       = "failed"
        job.error        = traceback.format_exc()
        job.completed_at = _now()
        _save()

        if job.notify_email:
            from app.email_notify import notify_failed
            notify_failed(job.notify_email, job.id)


def create_job(
    pdf_path: str,
    output_path: str,
    use_claude: bool = False,
    notify_email: str = "",
    user_id: int | None = None,
    original_stem: str = "",
) -> Job:
    job = Job(str(uuid.uuid4()), pdf_path, output_path, use_claude, notify_email, user_id, original_stem)
    _store(job)
    return job


def start_job(job: Job):
    """Enqueue job — uses Redis/RQ if available, otherwise a daemon thread."""
    if REDIS_URL:
        try:
            import redis
            from rq import Queue
            conn = redis.from_url(REDIS_URL)
            q    = Queue("ada", connection=conn)
            q.enqueue(_execute, job.id, job_timeout=3600)
            return
        except Exception:
            pass
    t = threading.Thread(target=_execute, args=(job.id,), daemon=True)
    t.start()
